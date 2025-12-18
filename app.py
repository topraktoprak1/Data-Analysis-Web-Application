from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import json
import os
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from functools import wraps

app = Flask(__name__)

# Enable CORS for React frontend (include Vite dev server origins)
CORS(app, supports_credentials=True, origins=[
    'http://localhost:3000', 'http://127.0.0.1:3000',
    'http://localhost:5173', 'http://127.0.0.1:5173'
])

# Security configuration
import secrets
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Session security
app.config['SESSION_COOKIE_SECURE'] = False  # Set True in production with HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = 28800  # 8 hours (28800 seconds)

# PostgreSQL configuration
# For local development, use: postgresql://username:password@localhost:5432/database_name
# For production, use environment variable
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 
    'postgresql://postgres:1234@localhost:8080/veri_analizi')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_size': 10,
    'pool_recycle': 3600,
    'pool_pre_ping': True,
}

db = SQLAlchemy(app)

# Cache for data to avoid reloading from database every time
_data_cache = {
    'data': None,
    'timestamp': None,
    'record_count': 0
}

def get_cache_key():
    """Get cache key based on database record count"""
    count = DatabaseRecord.query.count()
    return count

def clear_data_cache():
    """Clear the data cache"""
    global _data_cache
    _data_cache = {'data': None, 'timestamp': None, 'record_count': 0}

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)
    name = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), default='user')
    email = db.Column(db.String(120))
    first_name = db.Column(db.String(80))
    last_name = db.Column(db.String(80))
    profile_photo = db.Column(db.String(200), default='img/avatars/avatar1.jpeg')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class DatabaseRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    personel = db.Column(db.String(120), nullable=False)
    data = db.Column(db.Text, nullable=False)  # JSON string of record data
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class SavedFilter(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    filter_name = db.Column(db.String(200), nullable=False)
    filter_type = db.Column(db.String(50), nullable=False)  # 'database', 'pivot', 'graph'
    filter_config = db.Column(db.Text, nullable=False)  # JSON string of filter configuration
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)


# Initialize database and create admin user
def init_db():
    db.create_all()
    # Create admin if not exists
    admin = User.query.filter_by(username='admin').first()
    if not admin:
        admin = User(
            username='admin',
            password=generate_password_hash('admin123'),
            name='Administrator',
            role='admin'
        )
        db.session.add(admin)
        db.session.commit()
        print('✓ Admin user created: admin/admin123')
    else:
        print('✓ Database initialized')

# Login required decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            # For API endpoints, return JSON error instead of redirect
            if request.path.startswith('/api/'):
                return jsonify({'error': 'Unauthorized'}), 401
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

# Admin required decorator
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session or session.get('role') != 'admin':
            return jsonify({'error': 'Admin access required'}), 403
        return f(*args, **kwargs)
    return decorated_function

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Favorites file path
FAVORITES_FILE_PATH = os.path.join(os.path.dirname(__file__), 'favorite_reports.json')

# ============================================================================
# EXCEL FORMULA CALCULATION FUNCTIONS
# ============================================================================

# Cache for Excel reference data
_excel_cache = {
    'info_df': None,
    'hourly_rates_df': None,
    'summary_df': None,
    'file_path': None
}

def load_excel_reference_data(file_path=None):
    """Load Info, Hourly Rates, and Summary sheets from Excel file into cache"""
    global _excel_cache
    
    if file_path is None:
        # Try to find latest xlsb file in uploads
        upload_dir = app.config['UPLOAD_FOLDER']
        xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
        if not xlsb_files:
            return False
        xlsb_files.sort(reverse=True)
        file_path = os.path.join(upload_dir, xlsb_files[0])
    
    # Check if already cached
    if _excel_cache['file_path'] == file_path and _excel_cache['info_df'] is not None:
        return True
    
    try:
        # Read Info sheet
        df_info = pd.read_excel(file_path, sheet_name='Info', engine='pyxlsb')
        _excel_cache['info_df'] = df_info
        
        # Read Hourly Rates sheet (header on row 2)
        df_rates = pd.read_excel(file_path, sheet_name='Hourly Rates', engine='pyxlsb', header=1)
        _excel_cache['hourly_rates_df'] = df_rates
        
        # Try to read Summary sheet (may not exist in all files)
        try:
            df_summary = pd.read_excel(file_path, sheet_name='Summary', engine='pyxlsb')
            _excel_cache['summary_df'] = df_summary
        except:
            _excel_cache['summary_df'] = None
        
        _excel_cache['file_path'] = file_path
        print(f"✓ Loaded Excel reference data from {os.path.basename(file_path)}")
        return True
    except Exception as e:
        print(f"Error loading Excel reference data: {e}")
        return False

def xlookup(lookup_value, lookup_array, return_array, if_not_found=0):
    """Python implementation of Excel XLOOKUP function"""
    try:
        # Convert to pandas Series if needed
        if not isinstance(lookup_array, pd.Series):
            lookup_array = pd.Series(lookup_array)
        if not isinstance(return_array, pd.Series):
            return_array = pd.Series(return_array)
        
        # Handle different data types
        if pd.isna(lookup_value):
            return if_not_found
        
        # Try exact match first
        mask = lookup_array == lookup_value
        
        # If no exact match and both are strings, try normalized comparison
        if not mask.any() and isinstance(lookup_value, str):
            try:
                # Normalize by removing extra spaces and converting to uppercase
                normalized_lookup = ' '.join(str(lookup_value).strip().upper().split())
                normalized_array = lookup_array.astype(str).str.strip().str.upper().apply(lambda x: ' '.join(x.split()))
                mask = normalized_array == normalized_lookup
            except:
                pass
        
        if mask.any():
            idx = mask.idxmax()
            result = return_array.iloc[idx] if isinstance(return_array, pd.Series) else return_array[idx]
            return result if pd.notna(result) else if_not_found
        
        return if_not_found
    except Exception as e:
        print(f"XLOOKUP error: {e}")
        return if_not_found

def safe_float(value, default=0.0):
    """Safely convert value to float"""
    try:
        if pd.isna(value):
            return default
        return float(value)
    except:
        return default

def safe_str(value, default=''):
    """Safely convert value to string"""
    try:
        if pd.isna(value):
            return default
        return str(value).strip()
    except:
        return default

def calculate_auto_fields(record_data, file_path=None):
    """
    Calculate all auto-populated fields based on Excel formulas
    
    Args:
        record_data: Dictionary containing the manually entered record data
        file_path: Optional path to Excel file (uses latest if not provided)
    
    Returns:
        Tuple: (updated_record_data, list of fields that are N/A or empty)
    """
    # Load Excel reference data
    if not load_excel_reference_data(file_path):
        print("Warning: Could not load Excel reference data")
        return record_data, []
    
    info_df = _excel_cache['info_df']
    rates_df = _excel_cache['hourly_rates_df']
    summary_df = _excel_cache['summary_df']
    
    # Track which fields failed to calculate
    na_fields = []
    
    # Extract values from record_data (these are user inputs)
    person_id = safe_float(record_data.get('ID', 0))
    name_surname = safe_str(record_data.get('Name Surname', ''))
    discipline = safe_str(record_data.get('Discipline', ''))
    
    # Handle different variations of Week/Month field name
    week_month = safe_str(record_data.get('(Week /\nMonth)', '') or 
                         record_data.get('(Week / Month)', '') or 
                         record_data.get('Week / Month', '') or
                         record_data.get('Week/Month', ''))
    
    company = safe_str(record_data.get('Company', ''))
    projects_group = safe_str(record_data.get('Projects/Group', ''))
    scope = safe_str(record_data.get('Scope', ''))
    projects = safe_str(record_data.get('Projects', ''))
    nationality = safe_str(record_data.get('Nationality', ''))
    office_location = safe_str(record_data.get('Office Location', ''))
    
    # Handle different variations of TOTAL MH field name
    total_mh = safe_float(record_data.get('TOTAL\n MH', 0) or 
                         record_data.get('TOTAL MH', 0) or
                         record_data.get('Total MH', 0))
    
    kuzey_mh = safe_float(record_data.get('Kuzey MH', 0))
    kuzey_mh_person = safe_float(record_data.get('Kuzey MH-Person', 0))
    status = safe_str(record_data.get('Status', ''))
    
    print(f'DEBUG: total_mh = {total_mh}, week_month = {week_month}, currency will be calculated')
    
    # For İşveren fields that might be entered
    isveren_currency = safe_str(record_data.get('İşveren - Currency', ''))
    isveren_sozlesme_no = safe_str(record_data.get('İşveren- Sözleşme No', ''))
    isveren_hakedis_no = safe_str(record_data.get('İşveren- Hakediş No', ''))
    isveren_hakedis_donemi = safe_str(record_data.get('İşveren- Hakediş Dönemi', ''))
    isveren_hakedis_kapsam = safe_str(record_data.get('İşveren- Hakediş Kapsam', ''))
    isveren_mh_modifiye = safe_float(record_data.get('İşveren- MH-Modifiye', 0))
    
    # ========================================================================
    # FORMULA CALCULATIONS
    # ========================================================================
    
    # 1. North/South = XLOOKUP($G, Info!$N:$N, Info!$Q:$Q)
    # $G = Scope column (column G in DATABASE = Scope)
    # Info column N = index 13 (Scope), Info column Q = index 16 (North/South)
    north_south = xlookup(scope, info_df.iloc[:, 13], info_df.iloc[:, 16], '')
    # Set all possible variations of the field name
    record_data['North/\nSouth'] = north_south
    record_data['North/South'] = north_south
    record_data['North/ South'] = north_south  # With space after slash
    print(f'DEBUG: Calculated North/South = {north_south} for Scope = {scope}')
    
    # 2. Currency = IF(A=905264,"TL",XLOOKUP($A,'Hourly Rates'!$A:$A,'Hourly Rates'!$G:$G))
    if person_id == 905264:
        currency = 'TL'
    else:
        # Column G in Hourly Rates = index 6 (Currency 2)
        currency = xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 6], 'USD')
    record_data['Currency'] = currency
    
    print(f'DEBUG: person_id={person_id}, currency={currency}')
    
    # 3. Projects/Group = XLOOKUP($H, Info!$O:$O, Info!$P:$P)
    # $H = Projects column (column H in DATABASE = Projects)
    # Info column O = index 14 (Projects lookup), Info column P = index 15 (Projects/Group return)
    projects_group = xlookup(projects, info_df.iloc[:, 14], info_df.iloc[:, 15], '')
    record_data['Projects/Group'] = projects_group
    print(f'DEBUG: Calculated Projects/Group = {projects_group} for Projects = {projects}')
    
    # 16. AP-CB/Subcon = IF(ISNUMBER(SEARCH("AP-CB", E)), "AP-CB", "Subcon")
    # E = Company column
    ap_cb_subcon = 'AP-CB' if 'AP-CB' in company.upper() else 'Subcon'
    record_data['AP-CB /\nSubcon'] = ap_cb_subcon
    
    # 20. LS/Unit Rate = IF(OR((IFERROR(SEARCH("Lumpsum",G),0))>0,E="İ4",E="DEGENKOLB",E="Kilci Danışmanlık"),"Lumpsum","Unit Rate")
    # G = Scope, E = Company
    # Check if any condition is true:
    # 1. Scope contains "Lumpsum"
    # 2. Company is "İ4", "DEGENKOLB", or "Kilci Danışmanlık"
    scope_has_lumpsum = 'lumpsum' in scope.lower() if scope else False
    company_is_special = company in ['İ4', 'DEGENKOLB', 'Kilci Danışmanlık']
    
    if scope_has_lumpsum or company_is_special:
        ls_unit_rate = 'Lumpsum'
    else:
        ls_unit_rate = 'Unit Rate'
    
    record_data['LS/Unit Rate'] = ls_unit_rate
    print(f'DEBUG: LS/Unit Rate = {ls_unit_rate} (scope_has_lumpsum={scope_has_lumpsum}, company_is_special={company_is_special})')
    
    # 5. Hourly Base Rate = IF(AND(W="Subcon", AT="Unit Rate"),
    #    XLOOKUP($A,'Hourly Rates'!$A:$A,'Hourly Rates'!J:J),
    #    XLOOKUP($A,'Hourly Rates'!$A:$A,'Hourly Rates'!H:H))
    # W = AP-CB/Subcon, AT = LS/Unit Rate
    if ap_cb_subcon == 'Subcon' and ls_unit_rate == 'Unit Rate':
        # Column J = index 9 (Hourly Base Rates 3)
        hourly_base_rate = safe_float(xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 9], 0))
        print(f'DEBUG: Subcon+Unit Rate, looking up column J (index 9)')
    else:
        # Column H = index 7 (Hourly Base Rates 2)
        hourly_base_rate = safe_float(xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 7], 0))
        print(f'DEBUG: Not Subcon+Unit Rate (ap_cb_subcon={ap_cb_subcon}, ls_unit_rate={ls_unit_rate}), looking up column H (index 7)')
    
    record_data['Hourly Base Rate'] = hourly_base_rate
    print(f'DEBUG: hourly_base_rate = {hourly_base_rate}')
    
    # 6. Hourly Additional Rate = Complex nested IF
    # IF($AT="Lumpsum",0,IF($E="AP-CB",0,IF($E="AP-CB / pergel",0,
    # IF(P="USD",XLOOKUP($A,'Hourly Rates'!$A:$A,'Hourly Rates'!$L:$L),
    # IF(P="TL",XLOOKUP($A,'Hourly Rates'!$A:$A,'Hourly Rates'!$L:$L)*(XLOOKUP($D,Info!$U:$U,Info!$W:$W)))))))
    
    print(f'DEBUG: Calculating Hourly Additional Rate - ls_unit_rate={ls_unit_rate}, company={company}, currency={currency}')
    
    if ls_unit_rate == 'Lumpsum':
        hourly_additional_rate = 0
        print(f'DEBUG: Lumpsum detected, setting hourly_additional_rate = 0')
    elif company == 'AP-CB' or company == 'AP-CB / pergel':
        hourly_additional_rate = 0
        print(f'DEBUG: AP-CB company detected, setting hourly_additional_rate = 0')
    else:
        # Column L = index 11 (Hourly Additional Rates 2)
        additional_base = safe_float(xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 11], 0))
        print(f'DEBUG: additional_base from Hourly Rates = {additional_base}')
        
        if currency == 'USD':
            hourly_additional_rate = additional_base
        elif currency == 'TL':
            # Get TCMB USD/TRY rate from Info sheet
            # $D = week_month, Info!$U:$U = column 20 (Weeks/Month), Info!$W:$W = column 22 (TCMB USD/TRY)
            tcmb_rate = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 22], 1))
            hourly_additional_rate = additional_base * tcmb_rate
        else:
            hourly_additional_rate = 0
        print(f'DEBUG: Final hourly_additional_rate = {hourly_additional_rate}')
    
    record_data['Hourly Additional Rates'] = hourly_additional_rate
    
    print(f'DEBUG: Stored Hourly Additional Rates = {hourly_additional_rate}')
    
    # 3. Hourly Rate = S + V (Hourly Base Rate + Hourly Additional Rate)
    hourly_rate = hourly_base_rate + hourly_additional_rate
    record_data['Hourly\n Rate'] = hourly_rate
    record_data['Hourly Rate'] = hourly_rate
    
    # 4. Cost = Q * K (Hourly Rate * TOTAL MH)
    cost = hourly_rate * total_mh
    record_data['Cost'] = cost
    
    print(f'DEBUG: hourly_rate = {hourly_rate}, total_mh = {total_mh}, cost = {cost}')
    
    # 8. General Total Cost (USD) = IF($P="TL",$R/XLOOKUP($D,Info!$U:$U,Info!W:W),
    #    IF($P="EURO",$R*XLOOKUP($D,Info!$U:$U,Info!X:X),R))
    # P = Currency, R = Cost, D = week_month
    if currency == 'TL':
        tcmb_usd_try = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 22], 1))
        general_total_cost_usd = cost / tcmb_usd_try if tcmb_usd_try != 0 else 0
    elif currency == 'EURO':
        tcmb_eur_usd = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 23], 1))
        general_total_cost_usd = cost * tcmb_eur_usd
    else:
        general_total_cost_usd = cost
    record_data['General Total\n Cost (USD)'] = general_total_cost_usd
    record_data['General Total Cost (USD)'] = general_total_cost_usd
    
    print(f'DEBUG: currency = {currency}, cost = {cost}, general_total_cost_usd = {general_total_cost_usd}')
    
    # 9. Hourly Unit Rate (USD) = X / K (General Total Cost USD / TOTAL MH)
    hourly_unit_rate_usd = general_total_cost_usd / total_mh if total_mh != 0 else 0
    record_data['Hourly Unit Rate (USD)'] = hourly_unit_rate_usd
    record_data['Hourly Unit Rate (USD)'] = hourly_unit_rate_usd
    
    # Get NO-1, NO-2, NO-3, NO-10 values needed for İşveren calculations
    # 14. NO-1 = XLOOKUP($G,Info!$AU:$AU,Info!$AQ:$AQ,0)
    # Column AU = 46, Column AQ = 42
    print(f'DEBUG NO-1: Looking up scope="{scope}" in column AU (46)')
    print(f'DEBUG NO-1: Sample values in AU: {info_df.iloc[:5, 46].tolist()}')
    print(f'DEBUG NO-1: Sample values in AQ: {info_df.iloc[:5, 42].tolist()}')
    no_1 = xlookup(scope, info_df.iloc[:, 46], info_df.iloc[:, 42], 0)
    record_data['NO-1'] = no_1
    print(f'DEBUG: NO-1 lookup - scope="{scope}", result={no_1}')
    
    # 16. NO-2 = XLOOKUP($G,Info!$N:$N,Info!$L:$L)
    no_2 = xlookup(scope, info_df.iloc[:, 13], info_df.iloc[:, 11], '')
    record_data['NO-2'] = no_2
    print(f'DEBUG: NO-2 lookup - scope={scope}, result={no_2}')
    
    # 17. NO-3 = XLOOKUP($G,Info!$N:$N,Info!$M:$M)
    no_3 = xlookup(scope, info_df.iloc[:, 13], info_df.iloc[:, 12], '')
    record_data['NO-3'] = no_3
    print(f'DEBUG: NO-3 lookup - scope={scope}, result={no_3}')
    
    # 18. NO-10 = XLOOKUP($AN,Info!$J:$J,Info!$K:$K)
    # AN = NO-1
    no_10 = xlookup(no_1, info_df.iloc[:, 9], info_df.iloc[:, 10], '')
    record_data['NO-10'] = no_10
    print(f'DEBUG: NO-10 lookup - no_1={no_1}, result={no_10}')
    
    # 10. İşveren Hakediş Birim Fiyat (complex nested IF)
    # IF(OR(AQ="999-A", AQ="999-C", AQ="414-C", AN=313), Q,
    #    IF(OR(AN=312, AN=314, AN=316, AQ="360-T"), Q*1.02,
    #       IF(AQ="517-A", XLOOKUP(A, Info!AC:AC, Info!AH:AH),
    #          IFERROR(XLOOKUP(AN, Summary!C:C, Summary!AA:AA), 0)
    #          + IFERROR(XLOOKUP(AQ, Summary!C:C, Summary!AA:AA), 0))))
    # AQ = NO-2, AN = NO-1, Q = Hourly Rate, A = ID
    
    no_1_num = safe_float(no_1, 0)
    no_2_str = safe_str(no_2, '')
    
    if no_2_str in ['999-A', '999-C', '414-C'] or no_1_num == 313:
        isveren_hakedis_birim_fiyat = hourly_rate
    elif no_1_num in [312, 314, 316] or no_2_str == '360-T':
        isveren_hakedis_birim_fiyat = hourly_rate * 1.02
    elif no_2_str == '517-A':
        # XLOOKUP(A, Info!AC:AC, Info!AH:AH)
        # Column AC = index 28 (ID.1), Column AH = index 33
        isveren_hakedis_birim_fiyat = safe_float(xlookup(person_id, info_df.iloc[:, 28], info_df.iloc[:, 33], 0))
    else:
        # Use Summary sheet if available
        if summary_df is not None:
            # XLOOKUP(AN, Summary!C:C, Summary!AA:AA) + XLOOKUP(AQ, Summary!C:C, Summary!AA:AA)
            val1 = safe_float(xlookup(no_1, summary_df.iloc[:, 2], summary_df.iloc[:, 26], 0))
            val2 = safe_float(xlookup(no_2, summary_df.iloc[:, 2], summary_df.iloc[:, 26], 0))
            isveren_hakedis_birim_fiyat = val1 + val2
        else:
            isveren_hakedis_birim_fiyat = 0
    
    record_data['İşveren-Hakediş Birim Fiyat'] = isveren_hakedis_birim_fiyat
    
    # 11. İşveren-Hakediş(USD) = IF($L>0,(L*AA),AA*K)
    # L = Kuzey MH-Person, AA = İşveren Hakediş Birim Fiyat, K = TOTAL MH
    if kuzey_mh_person > 0:
        isveren_hakedis = kuzey_mh_person * isveren_hakedis_birim_fiyat
    else:
        isveren_hakedis = isveren_hakedis_birim_fiyat * total_mh
    record_data['İşveren- Hakediş'] = isveren_hakedis
    
    # 12. İşveren Hakediş (USD) = IF($Z="EURO",$AB*XLOOKUP($D,Info!$U:$U,Info!$X:$X),$AB)
    # Z = İşveren Currency, AB = İşveren-Hakediş, D = week_month
    if isveren_currency == 'EURO':
        eur_usd_rate = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 23], 1))
        isveren_hakedis_usd = isveren_hakedis * eur_usd_rate
    else:
        isveren_hakedis_usd = isveren_hakedis
    record_data['İşveren- Hakediş (USD)'] = isveren_hakedis_usd
    
    # 13. İşveren Hakediş Birim Fiyatı (USD) = IF($L>0,(AC/L),AC/K)
    # L = Kuzey MH-Person, AC = İşveren Hakediş (USD), K = TOTAL MH
    if kuzey_mh_person > 0:
        isveren_hakedis_birim_fiyat_usd = isveren_hakedis_usd / kuzey_mh_person if kuzey_mh_person != 0 else 0
    else:
        isveren_hakedis_birim_fiyat_usd = isveren_hakedis_usd / total_mh if total_mh != 0 else 0
    record_data['İşveren-Hakediş Birim Fiyat\n(USD)'] = isveren_hakedis_birim_fiyat_usd
    
    # 14. Control-1 = XLOOKUP(H,Info!O:O,Info!S:S)
    # H = Projects, Info O = column 14 (Projects), Info S = column 18 (Reporting)
    control_1 = xlookup(projects, info_df.iloc[:, 14], info_df.iloc[:, 18], '')
    record_data['Control-1'] = control_1
    print(f'DEBUG: Control-1 lookup - projects={projects}, result={control_1}')
    
    # 15. TM Liste = IFERROR(XLOOKUP(A,Info!BG:BG,Info!BI:BI),"")
    # Column BG = index 58, Column BI = index 60
    try:
        tm_liste = xlookup(person_id, info_df.iloc[:, 58], info_df.iloc[:, 60], '')
    except:
        tm_liste = ''
    record_data['TM Liste'] = tm_liste
    print(f'DEBUG: TM Liste lookup - person_id={person_id}, result={tm_liste}')
    
    # 16. TM KOD = XLOOKUP(H,Info!O:O,Info!R:R)
    # H = Projects, Info O = column 14, Info R = column 17 (TM KOD)
    tm_kod = xlookup(projects, info_df.iloc[:, 14], info_df.iloc[:, 17], '')
    record_data['TM Kod'] = tm_kod
    print(f'DEBUG: TM Kod lookup - projects={projects}, result={tm_kod}')
    
    # 17. Kontrol-1 = XLOOKUP(H,Info!AV:AV,Info!AQ:AQ)
    # H = Projects, Column AV = 47, Column AQ = 42
    print(f'DEBUG Kontrol-1: Looking up projects="{projects}" in column AV (47)')
    print(f'DEBUG Kontrol-1: Sample values in AV: {info_df.iloc[:5, 47].tolist()}')
    print(f'DEBUG Kontrol-1: Sample values in AQ: {info_df.iloc[:5, 42].tolist()}')
    kontrol_1 = xlookup(projects, info_df.iloc[:, 47], info_df.iloc[:, 42], '')
    record_data['Konrol-1'] = kontrol_1
    print(f'DEBUG: Kontrol-1 lookup - projects="{projects}", result={kontrol_1}')
    
    # 18. Kontrol-2 = AN=AO (NO-1 = Kontrol-1)
    kontrol_2 = no_1 == kontrol_1
    record_data['Knrtol-2'] = bool(kontrol_2)  # Convert to Python bool for JSON serialization
    
    # Check all calculated fields for N/A values
    fields_to_check = {
        'North/South': north_south,
        'Currency': currency,
        'Control-1': control_1,
        'TM Liste': tm_liste,
        'TM Kod': tm_kod,
        'Konrol-1': kontrol_1,
        'NO-1': no_1,
        'NO-2': no_2,
        'NO-3': no_3,
        'NO-10': no_10
    }
    
    for field_name, value in fields_to_check.items():
        # Check if value is empty, None, or 'N/A'
        if not value or value == '' or value == 'N/A' or (isinstance(value, float) and value == 0 and field_name.startswith('NO-')):
            na_fields.append(field_name)
            print(f'DEBUG: Field {field_name} has N/A or empty value')
    
    return record_data, na_fields

# Utility functions
def load_excel_data(file_path, user_filter=None):
    """Load Excel file and return DataFrame from DATABASE sheet with optional user filtering"""
    file_name = os.path.basename(file_path).lower()
    
    if file_name.endswith('.xlsb'):
        try:
            df = pd.read_excel(file_path, sheet_name='DATABASE', engine='pyxlsb')
        except:
            df = pd.read_excel(file_path, engine='pyxlsb')
    else:
        try:
            df = pd.read_excel(file_path, sheet_name='DATABASE')
        except:
            df = pd.read_excel(file_path)
    
    print(f"Loaded Excel with {len(df)} rows and {len(df.columns)} columns")
    
    # Clean column names - remove newlines and extra spaces
    df.columns = df.columns.str.replace('\n', ' ').str.replace('\r', ' ').str.replace('  ', ' ').str.strip()
    print(f"Cleaned column names")
    
    # Preserve date formats - convert datetime columns to string in dd/mmm/yyyy format
    for col in df.columns:
        # Check if column name contains date-related keywords (case insensitive, including parentheses)
        col_lower = str(col).lower() if col else ''
        if any(keyword in col_lower for keyword in ['week', 'month', 'date', 'tarih']):
            try:
                print(f"\nProcessing date column: '{col}'")
                print(f"Column dtype: {df[col].dtype}")
                print(f"First 5 raw values: {df[col].head(5).tolist()}")
                import sys
                sys.stdout.flush()  # Force output to appear immediately
                
                # If already datetime, just format it
                if df[col].dtype == 'datetime64[ns]':
                    # Manual formatting to ensure year is included
                    df[col] = df[col].apply(lambda x: x.strftime('%d/%b/%Y') if pd.notna(x) else '')
                    print(f"Converted datetime column directly")
                    print(f"First 5 after conversion: {df[col].head(5).tolist()}")
                else:
                    # Check if values are numbers (Excel date serial numbers)
                    first_val = df[col].dropna().iloc[0] if len(df[col].dropna()) > 0 else None
                    
                    if first_val is not None and isinstance(first_val, (int, float)) and first_val > 30000:
                        # These are Excel serial date numbers, convert them
                        print(f"Detected Excel serial numbers, converting...")
                        # Excel dates start from 1900-01-01 (serial 1)
                        # Convert Excel serial dates to pandas datetime
                        from datetime import datetime, timedelta
                        excel_start = datetime(1899, 12, 30)
                        date_series = df[col].apply(lambda x: excel_start + timedelta(days=float(x)) if pd.notna(x) and isinstance(x, (int, float)) else pd.NaT)
                        print(f"Successfully parsed {date_series.notna().sum()} out of {len(date_series)} values")
                        
                        # Format the dates manually to ensure year is included
                        def format_date(dt):
                            if pd.notna(dt):
                                try:
                                    # Manual string formatting
                                    day = dt.day
                                    month = dt.strftime('%b')
                                    year = dt.year
                                    formatted = f"{day:02d}/{month}/{year}"
                                    return formatted
                                except:
                                    return ''
                            return ''
                        
                        df[col] = date_series.apply(format_date)
                        print(f"First 5 after formatting: {df[col].head(5).tolist()}")
                        
                        # Extra debug: Check if year is actually there
                        sample_date = date_series.dropna().iloc[0] if len(date_series.dropna()) > 0 else None
                        if sample_date:
                            print(f"DEBUG: Sample datetime object - Day:{sample_date.day}, Month:{sample_date.month}, Year:{sample_date.year}")
                            print(f"DEBUG: Formatted sample: {format_date(sample_date)}")
                    else:
                        # Try standard datetime parsing
                        original_values = df[col].astype(str).copy()
                        date_series = pd.to_datetime(df[col], errors='coerce', format='mixed')
                        
                        success_count = date_series.notna().sum()
                        print(f"Successfully parsed {success_count} out of {len(date_series)} values")
                        
                        if success_count > 0:
                            formatted = date_series.dt.strftime('%d/%b/%Y')
                            df[col] = formatted.where(date_series.notna(), original_values)
                            print(f"First 5 after formatting: {df[col].head(5).tolist()}")
                        else:
                            print(f"No dates could be parsed, keeping original values")
                        
            except Exception as e:
                print(f"Error processing date column {col}: {e}")
                import traceback
                traceback.print_exc()
    
    # Filter by user if not admin
    if user_filter and 'PERSONEL' in df.columns:
        df = df[df['PERSONEL'].str.strip().str.upper() == user_filter.strip().upper()]
    
    return df

def get_data_from_db(user_filter=None):
    """Get data from database records with caching"""
    global _data_cache
    
    try:
        # DISABLE CACHE TEMPORARILY FOR DEBUGGING
        # Check if we can use cached data (only for no filter case)
        current_count = DatabaseRecord.query.count()
        
        print(f"DEBUG: Database has {current_count} records")
        
        # FORCE RELOAD - skip cache
        # if user_filter is None and _data_cache['data'] is not None and _data_cache['record_count'] == current_count:
        #     print(f"Using cached data ({current_count} records, {len(_data_cache['data'])} rows)")
        #     return _data_cache['data'].copy()
        
        # Load fresh data from database
        print(f"Loading data from database ({current_count} records)...")
        if user_filter:
            records = DatabaseRecord.query.filter_by(personel=user_filter).all()
            print(f"DEBUG: Filtered to {len(records)} records for user {user_filter}")
        else:
            records = DatabaseRecord.query.all()
            print(f"DEBUG: Loaded all {len(records)} records")
        
        if not records:
            print("DEBUG: No records found")
            return pd.DataFrame()
        
        records_list = []
        for record in records:
            try:
                record_dict = json.loads(record.data)
                # Debug: Check date format after loading from database
                if len(records_list) == 0 and '(Week / Month)' in record_dict:
                    print(f"DEBUG LOAD: First date value after DB load: {record_dict['(Week / Month)']}")
                records_list.append(record_dict)
            except:
                continue
        
        print(f"DEBUG: Parsed {len(records_list)} records into DataFrame")
        
        if records_list:
            # Create DataFrame without automatic date parsing
            df = pd.DataFrame(records_list, dtype=object)
            
            # Ensure date columns stay as strings (don't let pandas auto-convert)
            for col in df.columns:
                if 'week' in str(col).lower() or 'month' in str(col).lower() or 'date' in str(col).lower():
                    df[col] = df[col].astype(str)
            
            # DEBUG: Force print to console
            import sys
            print("\n" + "="*80, file=sys.stderr, flush=True)
            print(f"DEBUG DF: Created DataFrame with {len(df)} rows and {len(df.columns)} columns", file=sys.stderr, flush=True)
            
            # Debug: Check date column in DataFrame
            if '(Week / Month)' in df.columns:
                first_dates = df['(Week / Month)'].head(5).tolist()
                print(f"DEBUG DF: (Week / Month) first 5 values: {first_dates}", file=sys.stderr, flush=True)
                print(f"DEBUG DF: Sample date length: {len(first_dates[0]) if first_dates else 0}", file=sys.stderr, flush=True)
            print("="*80 + "\n", file=sys.stderr, flush=True)
            
            # Ensure PERSONEL column exists for filtering but don't expose it
            if 'Name Surname' in df.columns and 'PERSONEL' not in df.columns:
                df['PERSONEL'] = df['Name Surname']
            elif 'PERSONEL' in df.columns and 'Name Surname' not in df.columns:
                # If only PERSONEL exists, create Name Surname
                df['Name Surname'] = df['PERSONEL']
            
            # Cache the data if no filter
            if user_filter is None:
                _data_cache['data'] = df.copy()
                _data_cache['timestamp'] = datetime.now()
                _data_cache['record_count'] = current_count
                print(f"Data cached successfully")
            
            return df
        return pd.DataFrame()
    except Exception as e:
        print(f"Database load error: {str(e)}")
        return pd.DataFrame()

def get_combined_data(file_path=None, user_filter=None):
    """Get data from both Excel file and database"""
    dfs = []
    
    # Load from Excel if file exists
    if file_path and os.path.exists(file_path):
        try:
            df_excel = load_excel_data(file_path, user_filter)
            if not df_excel.empty:
                dfs.append(df_excel)
        except:
            pass
    
    # Load from database
    df_db = get_data_from_db(user_filter)
    if not df_db.empty:
        dfs.append(df_db)
    
    # Combine dataframes
    if dfs:
        combined_df = pd.concat(dfs, ignore_index=True)
        return combined_df
    return pd.DataFrame()

def add_calculated_columns(df):
    """Add KAR/ZARAR and BF KAR/ZARAR columns"""
    if df.empty:
        return df
    
    df = df.copy()
    
    # Check if columns already exist (cached data)
    if 'KAR/ZARAR' in df.columns and 'BF KAR/ZARAR' in df.columns:
        return df
    
    col_isveren = None
    col_general = None
    col_birim = None
    col_hourly = None
    
    for col in df.columns:
        col_clean = str(col).strip()
        if 'İşveren- Hakediş (USD)' in col_clean or 'İşveren-Hakediş (USD)' in col_clean:
            col_isveren = col
            print(f"Found İşveren column: {col}")
        elif 'General Total' in col_clean and 'Cost (USD)' in col_clean:
            col_general = col
            print(f"Found General Total column: {col}")
        elif 'İşveren-Hakediş Birim Fiyat' in col_clean and '(USD)' in col_clean:
            col_birim = col
            print(f"Found Birim Fiyat column: {col}")
        elif 'Hourly Unit Rate (USD)' in col_clean:
            col_hourly = col
            print(f"Found Hourly Rate column: {col}")
    
    if col_isveren and col_general:
        df['KAR/ZARAR'] = pd.to_numeric(df[col_isveren], errors='coerce') - pd.to_numeric(df[col_general], errors='coerce')
        print(f"Created KAR/ZARAR column with {df['KAR/ZARAR'].notna().sum()} valid values")
    else:
        print(f"WARNING: Could not create KAR/ZARAR - missing columns (İşveren: {col_isveren}, General: {col_general})")
    
    if col_birim and col_hourly:
        df['BF KAR/ZARAR'] = pd.to_numeric(df[col_birim], errors='coerce') - pd.to_numeric(df[col_hourly], errors='coerce')
        print(f"Created BF KAR/ZARAR column with {df['BF KAR/ZARAR'].notna().sum()} valid values")
    else:
        print(f"WARNING: Could not create BF KAR/ZARAR - missing columns (Birim: {col_birim}, Hourly: {col_hourly})")
    
    # Format (Week / Month) column if exists
    week_month_col = None
    for col in df.columns:
        if 'week' in str(col).lower() and 'month' in str(col).lower():
            week_month_col = col
            break
    
    if week_month_col:
        def format_week_month(val):
            try:
                if pd.isna(val):
                    return val
                dt = pd.to_datetime(val, errors='coerce')
                if pd.notna(dt):
                    return dt.strftime('%d/%b/%Y')
                return val
            except:
                return val
        
        df[week_month_col] = df[week_month_col].apply(format_week_month)
    
    return df

def load_favorites():
    """Load favorite reports from JSON file"""
    try:
        with open(FAVORITES_FILE_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, dict) and '_last_loaded' in data:
                return data
            else:
                return {'reports': data, '_last_loaded': None}
    except FileNotFoundError:
        return {'reports': {}, '_last_loaded': None}

def save_favorites(favorites, last_loaded=None):
    """Save favorite reports to JSON file"""
    data = {
        'reports': favorites,
        '_last_loaded': last_loaded
    }
    with open(FAVORITES_FILE_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)

# Routes
@app.route('/')
def home():
    if 'user' not in session:
        return redirect('/login.html')
    return redirect('/index.html')

@app.route('/index.html')
@login_required
def index():
    """Main dashboard page"""
    return render_template('index.html', user=session.get('user'), role=session.get('role'))

@app.route('/table.html')
@login_required
def table():
    """Pivot analysis page"""
    return render_template('table.html', user=session.get('user'), role=session.get('role'))

@app.route('/profile.html')
@login_required
def profile():
    """User profile page"""
    user = User.query.filter_by(username=session.get('user')).first()
    return render_template('profile.html', 
                         user=session.get('user'), 
                         name=user.name if user else session.get('user'),
                         role=session.get('role'))

@app.route('/login.html')
def login_page():
    """Login page"""
    return render_template('login.html')

@app.route('/register.html')
def register_page():
    """Registration page"""
    return render_template('register.html')

@app.route('/graphs.html')
@login_required
def graphs():
    """Graph analysis page"""
    return render_template('graphs.html', user=session.get('user'), role=session.get('role'))

@app.route('/admin.html')
@login_required
def admin_panel():
    """Admin panel - Admin only"""
    if session.get('role') != 'admin':
        return redirect('/index.html')
    return render_template('admin.html', user=session.get('user'), role=session.get('role'), name=session.get('name'))

# Simple rate limiting storage (in production, use Redis)
_login_attempts = {}

@app.route('/api/login', methods=['POST'])
def login():
    """Handle user login with rate limiting"""
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '')
    
    if not username or not password:
        return jsonify({'error': 'Username and password required'}), 400
    
    # Rate limiting check
    client_ip = request.remote_addr
    current_time = datetime.now().timestamp()
    
    if client_ip in _login_attempts:
        attempts, last_attempt = _login_attempts[client_ip]
        if current_time - last_attempt < 60:  # Within 1 minute
            if attempts >= 5:
                return jsonify({'error': 'Too many login attempts. Please try again in 1 minute.'}), 429
        else:
            _login_attempts[client_ip] = (0, current_time)
    
    # Check if user exists in database
    user = User.query.filter_by(username=username).first()
    if user and check_password_hash(user.password, password):
        # Reset failed attempts on successful login
        _login_attempts.pop(client_ip, None)
        
        session['user'] = username
        session['role'] = user.role
        session['name'] = user.name
        session.permanent = True
        return jsonify({
            'success': True,
            'role': user.role,
            'name': user.name
        })
    
    # Track failed attempts
    if client_ip in _login_attempts:
        attempts, _ = _login_attempts[client_ip]
        _login_attempts[client_ip] = (attempts + 1, current_time)
    else:
        _login_attempts[client_ip] = (1, current_time)
    
    return jsonify({'error': 'Invalid username or password'}), 401

@app.route('/api/register', methods=['POST'])
def register():
    """Handle user registration"""
    data = request.json
    username = data.get('username', '').strip()
    password = data.get('password', '')
    name = data.get('name', '').strip()
    
    if not username or not password or not name:
        return jsonify({'error': 'All fields required'}), 400
    
    # Password strength validation
    if len(password) < 8:
        return jsonify({'error': 'Password must be at least 8 characters'}), 400
    if not any(c.isupper() for c in password):
        return jsonify({'error': 'Password must contain at least one uppercase letter'}), 400
    if not any(c.isdigit() for c in password):
        return jsonify({'error': 'Password must contain at least one number'}), 400
    
    # Username validation
    if len(username) < 3:
        return jsonify({'error': 'Username must be at least 3 characters'}), 400
    if not username.isalnum():
        return jsonify({'error': 'Username must contain only letters and numbers'}), 400
    
    # Check if username exists
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400
    
    # Register new user (normal user role)
    new_user = User(
        username=username,
        password=generate_password_hash(password),
        name=name,
        role='user'
    )
    db.session.add(new_user)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Registration successful'})

@app.route('/api/profile', methods=['GET'])
@login_required
def get_profile():
    """Get current user profile"""
    user = User.query.filter_by(username=session.get('user')).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    return jsonify({
        'username': user.username,
        'name': user.name,
        'email': user.email or '',
        'first_name': user.first_name or '',
        'last_name': user.last_name or '',
        'profile_photo': user.profile_photo,
        'role': user.role
    })

@app.route('/api/profile', methods=['PUT'])
@login_required
def update_profile():
    """Update user profile"""
    try:
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        data = request.json
        
        # Update fields if provided
        if 'email' in data:
            user.email = data['email']
        if 'first_name' in data:
            user.first_name = data['first_name']
        if 'last_name' in data:
            user.last_name = data['last_name']
        if 'name' in data:
            user.name = data['name']
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Profile updated successfully'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/profile/photo', methods=['POST'])
@login_required
def upload_profile_photo():
    """Upload profile photo"""
    try:
        if 'photo' not in request.files:
            return jsonify({'error': 'No photo provided'}), 400
        
        file = request.files['photo']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_ext not in allowed_extensions:
            return jsonify({'error': 'Invalid file type. Use PNG, JPG, or GIF'}), 400
        
        user = User.query.filter_by(username=session.get('user')).first()
        if not user:
            return jsonify({'error': 'User not found'}), 404
        
        # Create profile photos directory if it doesn't exist
        photo_dir = os.path.join('static', 'img', 'profiles')
        os.makedirs(photo_dir, exist_ok=True)
        
        # Save photo with user ID in filename
        filename = f'user_{user.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.{file_ext}'
        filepath = os.path.join(photo_dir, filename)
        file.save(filepath)
        
        # Update user profile photo path
        user.profile_photo = f'img/profiles/{filename}'
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Photo uploaded successfully',
            'photo_url': f'/static/{user.profile_photo}'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/change-password', methods=['POST'])
@login_required
def change_password():
    """Change user password"""
    try:
        data = request.json
        current_password = data.get('current_password')
        new_password = data.get('new_password')
        
        # Validate input
        if not current_password or not new_password:
            return jsonify({
                'success': False,
                'message': 'Both current and new password are required'
            }), 400
        
        # Get current user
        user = User.query.get(session['user_id'])
        if not user:
            return jsonify({
                'success': False,
                'message': 'User not found'
            }), 404
        
        # Verify current password
        if not check_password_hash(user.password, current_password):
            return jsonify({
                'success': False,
                'message': 'Current password is incorrect'
            }), 401
        
        # Validate new password length
        if len(new_password) < 6:
            return jsonify({
                'success': False,
                'message': 'New password must be at least 6 characters long'
            }), 400
        
        # Check if new password is different from current
        if check_password_hash(user.password, new_password):
            return jsonify({
                'success': False,
                'message': 'New password must be different from current password'
            }), 400
        
        # Update password
        user.password = generate_password_hash(new_password)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Password changed successfully'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': f'Error changing password: {str(e)}'
        }), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    """Handle user logout"""
    session.clear()
    return jsonify({'success': True})

@app.route('/api/add-record', methods=['POST'])
@login_required
def add_record():
    """Add new database record - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        data = request.json
        record_data = data.get('record', {})
        
        # Get personel from either PERSONEL or Name Surname field
        personel = record_data.get('PERSONEL', '') or record_data.get('Name Surname', '')
        
        if not personel:
            return jsonify({'error': 'Name Surname (PERSONEL) field is required'}), 400
        
        # Ensure both fields are set for consistency
        if 'Name Surname' in record_data and 'PERSONEL' not in record_data:
            record_data['PERSONEL'] = record_data['Name Surname']
        
        # AUTO-CALCULATE fields based on Excel formulas
        # Get the latest Excel file path
        file_path = session.get('current_file')
        if not file_path or not os.path.exists(file_path):
            upload_dir = app.config['UPLOAD_FOLDER']
            xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
            if xlsb_files:
                xlsb_files.sort(reverse=True)
                file_path = os.path.join(upload_dir, xlsb_files[0])
        
        # Apply automatic calculations
        record_data, na_fields = calculate_auto_fields(record_data, file_path)
        
        # If there are N/A fields and no manual values provided, return them for user input
        if na_fields and not data.get('manual_values_provided'):
            return jsonify({
                'na_fields': na_fields,
                'message': f'Some fields could not be auto-calculated. Please provide values for: {", ".join(na_fields)}'
            }), 202  # 202 Accepted - needs more input
        
        # If manual values were provided in the second request, merge them
        if data.get('manual_values_provided') and data.get('manual_values'):
            manual_values = data.get('manual_values', {})
            for field, value in manual_values.items():
                if value and str(value).strip():  # Only set if not empty
                    record_data[field] = value
            print(f'DEBUG: Merged manual values: {manual_values}')
        
        # Create new record
        new_record = DatabaseRecord(
            personel=personel,
            data=json.dumps(record_data)
        )
        db.session.add(new_record)
        db.session.commit()
        
        # Clear cache to force reload
        clear_data_cache()
        
        return jsonify({'success': True, 'message': 'Record added successfully'})
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-records', methods=['GET'])
@login_required
def get_records():
    """Get database records filtered by user"""
    try:
        # Pagination parameters
        try:
            page = int(request.args.get('page', 1))
            per_page = int(request.args.get('per_page', 10))
        except:
            page = 1
            per_page = 10

        # Search parameter (search by personel / Name Surname)
        search = (request.args.get('search') or '').strip()

        # Build base query depending on role
        if session.get('role') == 'admin':
            base_query = DatabaseRecord.query
        else:
            base_query = DatabaseRecord.query.filter_by(personel=session.get('name'))

        # Apply search filter on the personel column
        if search:
            try:
                base_query = base_query.filter(DatabaseRecord.personel.ilike(f"%{search}%"))
            except Exception:
                # Fallback to simple filter (case-sensitive) if ilike not supported
                base_query = base_query.filter(DatabaseRecord.personel.contains(search))

        # Use SQLAlchemy pagination to avoid loading everything into memory
        pagination = base_query.order_by(DatabaseRecord.id.desc()).paginate(page=page, per_page=per_page, error_out=False)
        items = pagination.items

        records_list = []
        for record in items:
            record_dict = json.loads(record.data)
            record_dict['id'] = record.id
            
            # Remove duplicate field names (keep only the standard ones)
            if 'North/\nSouth' in record_dict:
                del record_dict['North/\nSouth']
            if 'North/ South' in record_dict:
                del record_dict['North/ South']
            
            records_list.append(record_dict)

        return jsonify({
            'success': True,
            'records': records_list,
            'page': page,
            'per_page': per_page,
            'pages': pagination.pages,
            'total': pagination.total
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/get-record/<int:record_id>', methods=['GET'])
@login_required
def get_record(record_id):
    """Get single database record by id"""
    try:
        record = DatabaseRecord.query.get(record_id)
        if not record:
            return jsonify({'error': 'Record not found'}), 404

        # Security: non-admins may only see their own records
        if session.get('role') != 'admin' and record.personel != session.get('name'):
            return jsonify({'error': 'Access denied'}), 403

        record_dict = json.loads(record.data)
        record_dict['id'] = record.id
        
        # Remove duplicate field names (keep only the standard ones)
        # Remove old variations of North/South
        if 'North/\nSouth' in record_dict:
            del record_dict['North/\nSouth']
        if 'North/ South' in record_dict:
            del record_dict['North/ South']
        
        return jsonify({'success': True, 'record': record_dict})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/update-record/<int:record_id>', methods=['PUT'])
@login_required
def update_record(record_id):
    """Update database record - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        record = DatabaseRecord.query.get(record_id)
        if not record:
            return jsonify({'error': 'Record not found'}), 404
        
        data = request.json
        record_data = data.get('record', {})
        
        # AUTO-RECALCULATE fields based on Excel formulas
        file_path = session.get('current_file')
        if not file_path or not os.path.exists(file_path):
            upload_dir = app.config['UPLOAD_FOLDER']
            xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
            if xlsb_files:
                xlsb_files.sort(reverse=True)
                file_path = os.path.join(upload_dir, xlsb_files[0])
        
        # Apply automatic calculations
        record_data = calculate_auto_fields(record_data, file_path)
        
        record.personel = record_data.get('PERSONEL', record.personel)
        record.data = json.dumps(record_data)
        record.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        # Clear cache to force reload
        clear_data_cache()
        
        return jsonify({'success': True, 'message': 'Record updated successfully'})
    except Exception as e:
        db.session.rollback()
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/delete-record/<int:record_id>', methods=['DELETE'])
@login_required
def delete_record(record_id):
    """Delete database record - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        record = DatabaseRecord.query.get(record_id)
        if not record:
            return jsonify({'error': 'Record not found'}), 404
        
        db.session.delete(record)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Record deleted successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-person-suggestions', methods=['GET'])
@login_required
def get_person_suggestions():
    """Get suggestions for a person based on their ID or historical data"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        person_id = request.args.get('id', '').strip()
        name = request.args.get('name', '').strip()
        
        if not person_id and not name:
            return jsonify({'error': 'ID or Name required'}), 400
        
        suggestions = {
            'found': False,
            'historical_data': {},
            'info_data': {},
            'warnings': [],
            'suggestions': []
        }
        
        # Try to get data from Info and Hourly Rates sheets
        file_path = session.get('current_file')
        if not file_path or not os.path.exists(file_path):
            upload_dir = app.config['UPLOAD_FOLDER']
            xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
            if xlsb_files:
                xlsb_files.sort(reverse=True)
                file_path = os.path.join(upload_dir, xlsb_files[0])
        
        if file_path and os.path.exists(file_path):
            load_excel_reference_data(file_path)
            info_df = _excel_cache.get('info_df')
            rates_df = _excel_cache.get('hourly_rates_df')
            
            if info_df is not None and person_id:
                # Search in Info sheet
                person_id_num = safe_float(person_id, 0)
                info_match = info_df[info_df.iloc[:, 0] == person_id_num]
                
                if not info_match.empty:
                    row = info_match.iloc[0]
                    suggestions['found'] = True
                    suggestions['info_data'] = {
                        'Name': safe_str(row.iloc[1]),
                        'Company': safe_str(row.iloc[2]),
                        'Nationality': safe_str(row.iloc[3]),
                        'Discipline': safe_str(row.iloc[6])
                    }
            
            if rates_df is not None and person_id:
                # Search in Hourly Rates sheet
                person_id_num = safe_float(person_id, 0)
                rates_match = rates_df[rates_df.iloc[:, 0] == person_id_num]
                
                if not rates_match.empty:
                    row = rates_match.iloc[0]
                    suggestions['found'] = True
                    suggestions['info_data']['Currency'] = safe_str(row.iloc[6])
                    suggestions['info_data']['Hourly_Base_Rate'] = safe_float(row.iloc[7], 0)
        
        # Get historical data from database
        if name or (suggestions['found'] and suggestions['info_data'].get('Name')):
            search_name = name or suggestions['info_data'].get('Name', '')
            historical_records = DatabaseRecord.query.filter_by(personel=search_name).all()
            
            if historical_records:
                # Analyze historical data
                total_mh_values = []
                hourly_rates = []
                
                for record in historical_records[-10:]:  # Last 10 records
                    try:
                        data = json.loads(record.data)
                        total_mh = safe_float(data.get('TOTAL MH', 0) or data.get('TOTAL\n MH', 0), 0)
                        hourly_rate = safe_float(data.get('Hourly Rate', 0) or data.get('Hourly\n Rate', 0), 0)
                        
                        if total_mh > 0:
                            total_mh_values.append(total_mh)
                        if hourly_rate > 0:
                            hourly_rates.append(hourly_rate)
                    except:
                        continue
                
                if total_mh_values:
                    avg_mh = sum(total_mh_values) / len(total_mh_values)
                    min_mh = min(total_mh_values)
                    max_mh = max(total_mh_values)
                    
                    suggestions['historical_data']['avg_total_mh'] = round(avg_mh, 2)
                    suggestions['historical_data']['min_total_mh'] = round(min_mh, 2)
                    suggestions['historical_data']['max_total_mh'] = round(max_mh, 2)
                    suggestions['historical_data']['record_count'] = len(historical_records)
                    
                    suggestions['suggestions'].append(
                        f"Historical average TOTAL MH: {round(avg_mh, 2)} hours (range: {round(min_mh, 2)} - {round(max_mh, 2)})"
                    )
                
                if hourly_rates:
                    avg_rate = sum(hourly_rates) / len(hourly_rates)
                    suggestions['historical_data']['avg_hourly_rate'] = round(avg_rate, 2)
                    suggestions['suggestions'].append(
                        f"Historical average Hourly Rate: ${round(avg_rate, 2)}"
                    )
        
        # Add general suggestions
        if suggestions['found']:
            suggestions['suggestions'].append("✓ Person found in Info/Hourly Rates sheets - data will be auto-filled")
        
        return jsonify({
            'success': True,
            'suggestions': suggestions
        })
    
    except Exception as e:
        print(f"Get person suggestions error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/validate-record', methods=['POST'])
@login_required
def validate_record():
    """Validate record data before saving and provide warnings"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        data = request.json
        record_data = data.get('record', {})
        
        warnings = []
        errors = []
        
        # Required field validation
        if not record_data.get('Name Surname') and not record_data.get('PERSONEL'):
            errors.append("Name Surname is required")
        
        if not record_data.get('ID'):
            warnings.append("ID is missing - calculations may not work correctly")
        
        # TOTAL MH validation
        total_mh = safe_float(record_data.get('TOTAL MH', 0) or record_data.get('TOTAL\n MH', 0), 0)
        
        if total_mh == 0:
            warnings.append("TOTAL MH is 0 - Cost will be 0")
        elif total_mh < 0:
            errors.append("TOTAL MH cannot be negative")
        elif total_mh > 500:
            warnings.append(f"TOTAL MH ({total_mh}) seems unusually high - please verify")
        elif total_mh < 1:
            warnings.append(f"TOTAL MH ({total_mh}) is less than 1 hour - please verify this is correct")
        
        # Week/Month validation
        week_month = safe_str(record_data.get('(Week /\nMonth)', '') or 
                             record_data.get('(Week / Month)', '') or 
                             record_data.get('Week / Month', ''))
        if not week_month:
            warnings.append("Week/Month is missing - currency conversion may not work")
        
        # Scope validation
        scope = safe_str(record_data.get('Scope', ''))
        if not scope:
            warnings.append("Scope is missing - North/South and other lookups will fail")
        
        return jsonify({
            'success': True,
            'valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings
        })
    
    except Exception as e:
        print(f"Validate record error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-person-info', methods=['GET'])
@login_required
def get_person_info():
    """Get person information from Info and Hourly Rates sheets"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        name = request.args.get('name', '').strip()
        if not name:
            return jsonify({'error': 'Name parameter required'}), 400
        
        # Get latest uploaded file
        file_path = session.get('current_file')
        if not file_path or not os.path.exists(file_path):
            # Try to find latest xlsb file in uploads
            upload_dir = app.config['UPLOAD_FOLDER']
            xlsb_files = [f for f in os.listdir(upload_dir) if f.endswith('.xlsb')]
            if not xlsb_files:
                return jsonify({'error': 'No Excel file found'}), 404
            xlsb_files.sort(reverse=True)
            file_path = os.path.join(upload_dir, xlsb_files[0])
        
        # Read Info sheet
        try:
            df_info = pd.read_excel(file_path, sheet_name='Info', engine='pyxlsb')
        except Exception as e:
            return jsonify({'error': f'Could not read Info sheet: {str(e)}'}), 500
        
        # Read Hourly Rates sheet (header is on row 2)
        try:
            df_rates = pd.read_excel(file_path, sheet_name='Hourly Rates', engine='pyxlsb', header=1)
        except Exception as e:
            return jsonify({'error': f'Could not read Hourly Rates sheet: {str(e)}'}), 500
        
        # Search in Info sheet by Name column
        info_match = df_info[df_info['Name'].str.strip().str.upper() == name.upper()]
        
        # Search in Hourly Rates by Name Surname column
        rates_match = df_rates[df_rates['Name Surname'].str.strip().str.upper() == name.upper()]
        
        result = {}
        
        # Extract data from Info sheet
        if not info_match.empty:
            row = info_match.iloc[0]
            # Do NOT pre-fill Company, Discipline, or Projects - these are manual inputs
            # result['Company'] = str(row.get('Company', '')) if pd.notna(row.get('Company')) else ''
            result['Nationality'] = str(row.get('Nationality', '')) if pd.notna(row.get('Nationality')) else ''
            # result['Discipline'] = str(row.get('Discipline', '')) if pd.notna(row.get('Discipline')) else ''
            result['Scope'] = str(row.get('Scope', '')) if pd.notna(row.get('Scope')) else ''
            # result['Projects'] = str(row.get('Projects', '')) if pd.notna(row.get('Projects')) else ''
            # Projects/Group will be auto-calculated from Projects, so don't pre-fill it
            # result['Projects/Group'] = str(row.get('Projects/Group', '')) if pd.notna(row.get('Projects/Group')) else ''
            result['North/South'] = str(row.get('North/South', '')) if pd.notna(row.get('North/South')) else ''
        
        # Extract data from Hourly Rates sheet
        if not rates_match.empty:
            row = rates_match.iloc[0]
            # Get hourly rates (trying different columns)
            for col in ['Hourly Base Rates 1', 'Hourly Base Rates 2', 'Hourly Base Rates 3']:
                if col in row and pd.notna(row[col]) and row[col] != 0:
                    result['Hourly Rate'] = float(row[col])
                    break
        
        if not result:
            return jsonify({'error': 'Person not found in Info or Hourly Rates sheets'}), 404
        
        return jsonify({
            'success': True,
            'data': result
        })
    
    except Exception as e:
        print(f"Get person info error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/clear-database', methods=['POST'])
@login_required
def clear_database():
    """Clear all database records - Admin only"""
    if session.get('role') != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    try:
        import sys
        print("\n" + "="*80, file=sys.stderr, flush=True)
        print("DEBUG: Starting database clear...", file=sys.stderr, flush=True)
        
        deleted = DatabaseRecord.query.delete()
        print(f"DEBUG: Deleted {deleted} records", file=sys.stderr, flush=True)
        
        db.session.commit()
        print("DEBUG: Database commit successful", file=sys.stderr, flush=True)
        
        # Clear cache
        clear_data_cache()
        
        # Clear session file reference
        session.pop('current_file', None)
        session.pop('data_shape', None)
        
        print("DEBUG: Clear database completed successfully", file=sys.stderr, flush=True)
        print("="*80 + "\n", file=sys.stderr, flush=True)
        
        return jsonify({
            'success': True, 
            'deleted': deleted,
            'message': f'Database cleared successfully! {deleted} records deleted.'
        })
    except Exception as e:
        import sys
        import traceback
        print("\n" + "="*80, file=sys.stderr, flush=True)
        print(f"ERROR clearing database: {str(e)}", file=sys.stderr, flush=True)
        traceback.print_exc(file=sys.stderr)
        print("="*80 + "\n", file=sys.stderr, flush=True)
        
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/forgot-password.html')
def forgot_password():
    """Forgot password page"""
    return render_template('forgot-password.html')

# API Routes
@app.route('/api/check-session', methods=['GET'])
@login_required
def check_session():
    """Check if file is loaded in session or data exists in database"""
    try:
        # Get user filter (None for admin, name for regular users)
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Get combined data from file and database
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        
        if df.empty:
            return jsonify({'hasData': False})
        
        # Add calculated columns (will use cached if available)
        df_with_calc = add_calculated_columns(df)
        
        # Get all columns including calculated ones
        all_columns = df_with_calc.columns.tolist()
        
        # Remove PERSONEL from column list if Name Surname exists (it's internal only)
        if 'Name Surname' in all_columns and 'PERSONEL' in all_columns:
            all_columns.remove('PERSONEL')
        
        # Limit preview to first 5 rows to prevent page collapse
        df_preview = df_with_calc.head(5).fillna('')
        data_json = df_preview.to_dict('records')
        
        # Get filter columns for ALL categorical columns
        filter_cols = []
        # Skip these columns from filters (PERSONEL is internal, but Name Surname is allowed)
        skip_cols = ['PERSONEL', 'id', 'created_at', 'updated_at']
        
        # Always add Name Surname filter first (important for filtering by person)
        if 'Name Surname' in df_with_calc.columns:
            name_values = sorted([str(v) for v in df_with_calc['Name Surname'].dropna().unique()])
            if name_values:
                filter_cols.append({
                    'name': 'Name Surname',
                    'values': name_values
                })
        
        for col in df_with_calc.columns:
            # Skip numeric columns and internal columns
            if col in skip_cols or col == 'Name Surname':  # Skip Name Surname as it's already added
                continue
                
            # Check if column is categorical (not purely numeric)
            try:
                # Try to detect categorical columns
                unique_count = df_with_calc[col].nunique()
                total_count = len(df_with_calc[col].dropna())
                
                # If column has reasonable number of unique values, add to filters
                # Max 200 unique values to accommodate other columns
                if unique_count > 0 and unique_count <= 200:
                    # Check if it's not a purely numeric column with many values
                    non_null_values = df_with_calc[col].dropna()
                    if len(non_null_values) > 0:
                        # Debug BEFORE creating unique values
                        if 'week' in col.lower() or 'month' in col.lower():
                            import sys
                            print("\n" + "="*80, file=sys.stderr, flush=True)
                            print(f"DEBUG RAW DF: Column '{col}'", file=sys.stderr, flush=True)
                            print(f"DEBUG RAW DF: First 10 raw values from df: {df_with_calc[col].head(10).tolist()}", file=sys.stderr, flush=True)
                            print(f"DEBUG RAW DF: Dtype: {df_with_calc[col].dtype}", file=sys.stderr, flush=True)
                        
                        # Add to filters
                        unique_values = sorted([str(v) for v in df_with_calc[col].dropna().unique()])
                        
                        # Debug date column - FORCE OUTPUT
                        if 'week' in col.lower() or 'month' in col.lower():
                            print(f"DEBUG FILTER: Total unique values: {len(unique_values)}", file=sys.stderr, flush=True)
                            print(f"DEBUG FILTER: First 5: {unique_values[:5]}", file=sys.stderr, flush=True)
                            print(f"DEBUG FILTER: Last 5: {unique_values[-5:]}", file=sys.stderr, flush=True)
                            if unique_values:
                                print(f"DEBUG FILTER: Sample length: {len(unique_values[0])}", file=sys.stderr, flush=True)
                            print("="*80 + "\n", file=sys.stderr, flush=True)
                        
                        filter_cols.append({
                            'name': col,
                            'values': unique_values
                        })
            except:
                continue
        
        return jsonify({
            'hasData': True,
            'columns': all_columns,
            'shape': df_with_calc.shape,
            'data': data_json,  # Preview only (first 5 rows)
            'filter_columns': filter_cols
        })
    except Exception as e:
        print(f"Check session error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'hasData': False, 'error': str(e)})

@app.route('/api/get-input-fields', methods=['GET'])
@login_required
def get_input_fields():
    """Get input fields and calculated fields from database schema"""
    try:
        # Get user filter (None for admin, name for regular users)
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Get combined data from file and database
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        
        # Define auto-calculated fields (these should NOT appear in the add record form)
        auto_calculated_fields = [
            'North/\nSouth',
            'North/South',
            'North/ South',
            'Currency',
            'Projects/Group',
            'Hourly\n Rate',
            'Hourly Rate',
            'Cost',
            'Hourly Base Rate',
            'Hourly Additional Rates',
            'AP-CB /\nSubcon',
            'AP-CB/Subcon',
            'AP-CB / Subcon',
            'LS/Unit Rate',
            'General Total\n Cost (USD)',
            'General Total Cost (USD)',
            'Hourly Unit Rate (USD)',
            'İşveren-Hakediş Birim Fiyat',
            'İşveren- Hakediş',
            'İşveren- Hakediş (USD)',
            'İşveren-Hakediş Birim Fiyat\n(USD)',
            'İşveren-Hakediş Birim Fiyat (USD)',
            'Control-1',
            'Control-1\n TM Liste',
            'Control-1\n TM Kod',
            'TM Liste',
            'TM Kod',
            'NO-1',
            'Kontrol-1',
            'Kontrol-2',
            'Konrol-1',
            'Knrtol-2',
            'NO-2',
            'NO-3',
            'NO-10',
            'KAR/ZARAR',
            'BF KAR/ZARAR'
        ]
        
        if df.empty:
            # Return default fields if no data exists
            default_input_fields = [
                'ID',
                'Name Surname',
                'Discipline',
                '(Week /\nMonth)',
                'Company',
                'Projects/Group',
                'Scope',
                'Projects',
                'Nationality',
                'Office Location',
                'TOTAL\n MH',
                'Kuzey MH',
                'Kuzey MH-Person',
                'Status',
                'PP',
                'Actual Value',
                'İşveren - Currency',
                'İşveren- Sözleşme No',
                'İşveren- Hakediş No',
                'İşveren- Hakediş Dönemi',
                'İşveren- Hakediş Kapsam',
                'İşveren- MH-Modifiye',
                'İşveren- Actual (USD)'
            ]
            
            return jsonify({
                'success': True,
                'input_fields': default_input_fields,
                'calculated_fields': auto_calculated_fields
            })
        
        # Get all columns from data
        all_columns = df.columns.tolist()
        
        # Input fields are all columns except auto-calculated ones
        input_fields = [col for col in all_columns if col not in auto_calculated_fields]
        
        # Remove internal PERSONEL if Name Surname exists
        if 'Name Surname' in input_fields and 'PERSONEL' in input_fields:
            input_fields.remove('PERSONEL')
        
        return jsonify({
            'success': True,
            'input_fields': input_fields,
            'calculated_fields': auto_calculated_fields
        })
    
    except Exception as e:
        print(f"Get input fields error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': str(e),
            'input_fields': ['ID', 'Name Surname', 'Discipline', '(Week /\nMonth)', 'Company', 'Projects/Group', 
                            'Scope', 'Projects', 'Nationality', 'Office Location', 'TOTAL\n MH', 'Status'],
            'calculated_fields': []
        }), 500

@app.route('/api/upload', methods=['POST'])
@login_required
def upload_file():
    """Handle file upload and process Excel data"""
    try:
        # Clear cache on new upload
        clear_data_cache()
        
        # Only admin can upload files
        if session.get('role') != 'admin':
            return jsonify({'error': 'Only admin can upload files'}), 403
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file:
            return jsonify({'error': 'Invalid file'}), 400
            
        # Check file extension
        allowed_extensions = {'.xlsx', '.xls', '.xlsb'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'Invalid file type. Allowed: {", ".join(allowed_extensions)}'}), 400
        
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{timestamp}_{filename}'
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Save file
        print(f"Saving file to: {filepath}")
        file.save(filepath)
        print(f"File saved successfully")
        
        # Load and process data quickly (don't apply calculated columns yet for speed)
        print(f"Loading Excel data...")
        df = load_excel_data(filepath)
        print(f"Data loaded: {df.shape[0]} rows, {df.shape[1]} columns")
        
        # Save data to database permanently
        print(f"Saving data to database...")
        saved_count = 0
        skipped_count = 0
        
        # Clear existing database records (optional - remove if you want to keep old data)
        # DatabaseRecord.query.delete()
        
        for idx, row in df.iterrows():
            try:
                # Convert row to dictionary
                row_dict = row.to_dict()
                
                # Get personel name from either PERSONEL or Name Surname column
                personel = row_dict.get('PERSONEL') or row_dict.get('Name Surname', '')
                
                if not personel:
                    skipped_count += 1
                    continue
                
                # Ensure both fields exist for consistency
                if 'Name Surname' in row_dict and 'PERSONEL' not in row_dict:
                    row_dict['PERSONEL'] = row_dict['Name Surname']
                elif 'PERSONEL' in row_dict and 'Name Surname' not in row_dict:
                    row_dict['Name Surname'] = row_dict['PERSONEL']
                
                # Convert any NaN values to empty strings
                for key, value in row_dict.items():
                    if pd.isna(value):
                        row_dict[key] = ''
                
                # Debug: Check date format before saving to database
                if '(Week / Month)' in row_dict:
                    if saved_count == 0:  # Only log first record
                        print(f"DEBUG SAVE: First date value before DB save: {row_dict['(Week / Month)']}")
                
                # Create database record
                new_record = DatabaseRecord(
                    personel=str(personel),
                    data=json.dumps(row_dict)
                )
                db.session.add(new_record)
                saved_count += 1
                
                # Commit in batches for better performance
                if saved_count % 100 == 0:
                    db.session.commit()
                    print(f"Saved {saved_count} records...")
                    
            except Exception as e:
                print(f"Error saving row {idx}: {str(e)}")
                skipped_count += 1
                continue
        
        # Final commit
        db.session.commit()
        print(f"Database save complete: {saved_count} saved, {skipped_count} skipped")
        
        # Store in session
        session['current_file'] = filepath
        session['data_shape'] = df.shape
        
        # Get important filter columns only (for speed)
        filter_cols = []
        important_cols = ['PERSONEL', 'Name Surname', 'Company', 'Projects', 'Status', 'Discipline']
        for col in important_cols:
            if col in df.columns:
                unique_values = df[col].dropna().unique()
                if len(unique_values) <= 50:
                    filter_cols.append({
                        'name': col,
                        'values': [str(v) for v in unique_values[:50]]
                    })
        
        # Apply calculated columns after basic info is ready
        df = add_calculated_columns(df)
        
        # Convert to JSON-friendly format (full dataset)
        df_clean = df.fillna('')
        data_json = df_clean.to_dict('records')
        
        print(f"Upload successful, returning response")
        return jsonify({
            'success': True,
            'shape': df.shape,
            'columns': df.columns.tolist(),
            'filter_columns': filter_cols,
            'data': data_json,
            'saved_to_db': saved_count,
            'skipped': skipped_count,
            'message': f'File uploaded successfully! {df.shape[0]} rows, {df.shape[1]} columns. Saved {saved_count} records to database.'
        })
    
    except Exception as e:
        print(f"Upload error: {str(e)}")
        import traceback
        traceback.print_exc()
        db.session.rollback()  # Rollback any pending transactions
        return jsonify({'error': str(e)}), 500


def find_column(df, *possible_names):
    """Find which variation of a column name exists in the DataFrame"""
    for name in possible_names:
        if name in df.columns:
            return name
    return None

def set_if_empty(df, idx, col_name, value, debug=False):
    """Set value only if the cell is empty and column exists"""
    if col_name and col_name in df.columns:
        current_val = df.at[idx, col_name]
        
        # Check if cell is truly empty:
        # 1. NaN (most common for numeric empty cells)
        # 2. None
        # 3. Empty string or whitespace-only string
        # Note: We do NOT treat 0 as empty - it's a valid calculated value
        is_empty = False
        
        if pd.isna(current_val):
            is_empty = True
        elif current_val is None:
            is_empty = True
        elif isinstance(current_val, str) and current_val.strip() == '':
            is_empty = True
        
        if debug and idx == 0:  # Debug first row only
            print(f"  {col_name}: current={current_val} (type={type(current_val).__name__}), is_empty={is_empty}, will_set={value}")
        
        if is_empty:
            # If setting a string value to a numeric column, convert column to object dtype
            if isinstance(value, str) and df[col_name].dtype in ['float64', 'int64']:
                df[col_name] = df[col_name].astype('object')
            
            df.at[idx, col_name] = value
            return True
    return False

def fill_empty_cells_with_formulas(df, info_df, rates_df, summary_df):
    """
    Fill empty cells in DATABASE sheet based on Excel formulas.
    
    This function processes each row and fills empty cells using the same
    formulas that would be used in Excel, leveraging XLOOKUP and other functions.
    
    Args:
        df: DataFrame from DATABASE sheet (to be filled)
        info_df: DataFrame from Info sheet
        rates_df: DataFrame from Hourly Rates sheet  
        summary_df: DataFrame from Summary sheet
        
    Returns:
        DataFrame with empty cells filled
    """
    print(f"Starting to fill empty cells for {len(df)} rows...")
    print(f"Columns in DataFrame: {df.columns.tolist()}")
    
    # Create a copy to avoid modifying original
    result_df = df.copy()
    
    # Find which column variations exist
    col_north_south = find_column(result_df, 'North/South', 'North/\nSouth', 'North/ South')
    col_currency = find_column(result_df, 'Currency')
    col_ap_cb_subcon = find_column(result_df, 'AP-CB / \nSubcon')
    col_ls_unit_rate = find_column(result_df, 'LS/Unit Rate')
    col_hourly_base_rate = find_column(result_df, 'Hourly Base Rate')
    col_hourly_additional_rate = find_column(result_df, 'Hourly Additional Rates', 'Hourly Additional Rate')
    col_hourly_rate = find_column(result_df, 'Hourly Rate', 'Hourly\n Rate')
    col_cost = find_column(result_df, 'Cost')
    col_general_total_cost = find_column(result_df, 'General Total Cost (USD)', 'General Total\n Cost (USD)')
    col_hourly_unit_rate_usd = find_column(result_df, 'Hourly Unit Rate (USD)')
    col_no_1 = find_column(result_df, 'NO-1')
    col_no_2 = find_column(result_df, 'NO-2')
    col_no_3 = find_column(result_df, 'NO-3')
    col_no_10 = find_column(result_df, 'NO-10')
    col_isveren_birim_fiyat = find_column(result_df, 'İşveren-Hakediş Birim Fiyat')
    col_isveren_hakedis = find_column(result_df, 'İşveren- Hakediş')
    col_isveren_hakedis_usd = find_column(result_df, 'İşveren- Hakediş (USD)')
    col_isveren_birim_fiyat_usd = find_column(result_df, 'İşveren-Hakediş Birim Fiyat\n(USD)', 'İşveren-Hakediş Birim Fiyat (USD)')
    col_control_1 = find_column(result_df, 'Control-1')
    col_tm_liste = find_column(result_df, 'TM Liste')
    col_tm_kod = find_column(result_df, 'TM Kod')
    col_kontrol_1 = find_column(result_df, 'Konrol-1', 'Kontrol-1')
    col_kontrol_2 = find_column(result_df, 'Knrtol-2', 'Kontrol-2')
    
    print(f"Found columns - Currency: {col_currency}, Hourly Rate: {col_hourly_rate}, Cost: {col_cost}")
    print(f"Found columns - NO-3: {col_no_3}, Kontrol-1: {col_kontrol_1}, Kontrol-2: {col_kontrol_2}")
    print(f"All columns in DataFrame: {result_df.columns.tolist()}")
    
    # Process each row
    for idx, row in result_df.iterrows():
        if idx % 100 == 0:
            print(f"Processing row {idx}...")
        
        # Extract base values (these should already exist in the file)
        person_id = safe_float(row.get('ID', 0))
        scope = safe_str(row.get('Scope', ''))
        company = safe_str(row.get('Company', ''))
        projects = safe_str(row.get('Projects', ''))
        
        # Handle different variations of Week/Month field
        week_month = safe_str(row.get('(Week / Month)', '') or 
                             row.get('(Week /\nMonth)', '') or
                             row.get('Week / Month', '') or
                             row.get('Week/Month', ''))
        
        # Handle different variations of TOTAL MH field
        total_mh = safe_float(row.get('TOTAL MH', 0) or 
                             row.get('TOTAL\n MH', 0) or
                             row.get('Total MH', 0))
        
        kuzey_mh_person = safe_float(row.get('Kuzey MH-Person', 0))
        
        # Get İşveren fields if they exist
        isveren_currency = safe_str(row.get('İşveren - Currency', '') or row.get('İşveren-Currency', ''))
        
        # ============================================================
        # FORMULA 1: North/South = XLOOKUP($G,Info!$N:$N,Info!$Q:$Q)
        # ============================================================
        if col_north_south:
            north_south = xlookup(scope, info_df.iloc[:, 13], info_df.iloc[:, 16], '')
            set_if_empty(result_df, idx, col_north_south, north_south)
        
        # ============================================================
        # FORMULA 2: Currency = IF(A=905264,"TL",XLOOKUP($A,'Hourly Rates'!$A:$A,'Hourly Rates'!$G:$G))
        # ============================================================
        if col_currency:
            if person_id == 905264:
                currency = 'TL'
            else:
                currency = xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 6], 'USD')
            set_if_empty(result_df, idx, col_currency, currency)
            currency = result_df.at[idx, col_currency]  # Get actual value (might not have been set)
        else:
            currency = 'USD'  # Default
        
        # ============================================================
        # FORMULA 3: AP-CB/Subcon = IF(ISNUMBER(SEARCH("AP-CB", E)), "AP-CB", "Subcon")
        # ============================================================
       
        if col_ap_cb_subcon:
           ap_cb_subcon = 'AP-CB' if 'AP-CB' in company.upper() else 'Subcon'
           set_if_empty(result_df, idx, col_ap_cb_subcon, ap_cb_subcon)

    # Get actual value (might not have been set if cell already had data)
           ap_cb_subcon = result_df.at[idx, col_ap_cb_subcon]
        else:
    # If column doesn't exist, calculate anyway for use in other formulas
            ap_cb_subcon = 'AP-CB' if 'AP-CB' in company.upper() else 'Subcon'

        
        # ============================================================
        # FORMULA 4: LS/Unit Rate
        # ============================================================
        if col_ls_unit_rate:
            scope_has_lumpsum = 'lumpsum' in scope.lower() if scope else False
            company_is_special = company in ['İ4', 'DEGENKOLB', 'Kilci Danışmanlık']
            ls_unit_rate = 'Lumpsum' if (scope_has_lumpsum or company_is_special) else 'Unit Rate'
            set_if_empty(result_df, idx, col_ls_unit_rate, ls_unit_rate)
            ls_unit_rate = result_df.at[idx, col_ls_unit_rate]  # Get actual value (might not have been set)
        else:
            # If column doesn't exist, calculate anyway for use in other formulas
            scope_has_lumpsum = 'lumpsum' in scope.lower() if scope else False
            company_is_special = company in ['İ4', 'DEGENKOLB', 'Kilci Danışmanlık']
            ls_unit_rate = 'Lumpsum' if (scope_has_lumpsum or company_is_special) else 'Unit Rate'
        
        # ============================================================
        # FORMULA 5: Hourly Base Rate
        # ============================================================
        if ap_cb_subcon == 'Subcon' and ls_unit_rate == 'Unit Rate':
            hourly_base_rate = safe_float(xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 9], 0))
        else:
            hourly_base_rate = safe_float(xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 7], 0))
        
        if idx == 0:
            print(f"\n=== DEBUG ROW 0 - Hourly Base Rate ===")
            print(f"  person_id: {person_id}")
            print(f"  ap_cb_subcon: {ap_cb_subcon}, ls_unit_rate: {ls_unit_rate}")
            print(f"  rates_df shape: {rates_df.shape}")
            print(f"  rates_df columns: {rates_df.columns.tolist()}")
            print(f"  First 3 IDs in rates: {rates_df.iloc[:3, 0].tolist()}")
            print(f"  Column 7 (index 7) first 3 values: {rates_df.iloc[:3, 7].tolist()}")
            print(f"  Column 9 (index 9) first 3 values: {rates_df.iloc[:3, 9].tolist()}")
            print(f"  Calculated hourly_base_rate: {hourly_base_rate}")
            print("=" * 50)
        
        if col_hourly_base_rate:
            set_if_empty(result_df, idx, col_hourly_base_rate, hourly_base_rate, debug=(idx==0))
            hourly_base_rate = safe_float(result_df.at[idx, col_hourly_base_rate])
            
            if idx == 0:
                print(f"  Final hourly_base_rate after set: {hourly_base_rate}\n")
        
        # ============================================================
        # FORMULA 6: Hourly Additional Rate
        # ============================================================
        if ls_unit_rate == 'Lumpsum':
            hourly_additional_rate = 0
        elif company == 'AP-CB' or company == 'AP-CB / pergel':
            hourly_additional_rate = 0
        else:
            additional_base = safe_float(xlookup(person_id, rates_df.iloc[:, 0], rates_df.iloc[:, 11], 0))
            if currency == 'USD':
                hourly_additional_rate = additional_base
            elif currency == 'TL':
                tcmb_rate = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 22], 1))
                hourly_additional_rate = additional_base * tcmb_rate
            else:
                hourly_additional_rate = 0
        
        if col_hourly_additional_rate:
            set_if_empty(result_df, idx, col_hourly_additional_rate, hourly_additional_rate)
            hourly_additional_rate = safe_float(result_df.at[idx, col_hourly_additional_rate])
        
        # ============================================================
        # FORMULA 7: Hourly Rate = S + V
        # ============================================================
        hourly_rate = hourly_base_rate + hourly_additional_rate
        if col_hourly_rate:
            set_if_empty(result_df, idx, col_hourly_rate, hourly_rate)
            hourly_rate = safe_float(result_df.at[idx, col_hourly_rate])
        
        # ============================================================
        # FORMULA 8: Cost = Q * K
        # ============================================================
        cost = hourly_rate * total_mh
        if col_cost:
            set_if_empty(result_df, idx, col_cost, cost)
            cost = safe_float(result_df.at[idx, col_cost])
        
        # ============================================================
        # FORMULA 9: General Total Cost (USD)
        # ============================================================
        if currency == "TL":
            tcmb_rate = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 22], 1))
            general_total_cost_usd = cost / tcmb_rate if tcmb_rate != 0 else 0
        elif currency == "EURO":
            tcmb_eur_usd = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 23], 1))
            general_total_cost_usd = cost * tcmb_eur_usd
        else:
            general_total_cost_usd = cost
        
        if col_general_total_cost:
            set_if_empty(result_df, idx, col_general_total_cost, general_total_cost_usd)
            general_total_cost_usd = safe_float(result_df.at[idx, col_general_total_cost])

        
        # ============================================================
        # FORMULA 10: Hourly Unit Rate (USD)
        # ============================================================
        hourly_unit_rate_usd = general_total_cost_usd / total_mh if total_mh != 0 else 0
        if col_hourly_unit_rate_usd:
            set_if_empty(result_df, idx, col_hourly_unit_rate_usd, hourly_unit_rate_usd)
            hourly_unit_rate_usd = 0
        
        # ============================================================
        # Get NO-1, NO-2, NO-3, NO-10 for İşveren calculations
        # ============================================================
        no_1 = xlookup(scope, info_df.iloc[:, 46], info_df.iloc[:, 42], 0)
        if col_no_1:
            set_if_empty(result_df, idx, col_no_1, no_1)
            no_1 = result_df.at[idx, col_no_1]
        
        no_2 = xlookup(scope, info_df.iloc[:, 13], info_df.iloc[:, 11], '')
        if col_no_2:
            set_if_empty(result_df, idx, col_no_2, no_2)
            no_2 = result_df.at[idx, col_no_2]
        
        no_3 = xlookup(scope, info_df.iloc[:, 13], info_df.iloc[:, 12], '')
        
        if idx == 0:
            print(f"\n=== DEBUG ROW 0 - NO-3 ===")
            print(f"  col_no_3 column name: '{col_no_3}'")
            print(f"  scope: {scope}")
            print(f"  Calculated no_3: {no_3} (type: {type(no_3).__name__})")
            if col_no_3:
                print(f"  Current value at [0, '{col_no_3}']: {result_df.at[0, col_no_3]}")
                print(f"  Column dtype: {result_df[col_no_3].dtype}")
            print("=" * 50)
        
        if col_no_3:
            # If no_3 is empty string and column is numeric, convert to NaN or 0
            if no_3 == '' and result_df[col_no_3].dtype in ['float64', 'int64']:
                no_3 = 0  # or use np.nan if you want NaN
            set_if_empty(result_df, idx, col_no_3, no_3)
            no_3 = result_df.at[idx, col_no_3]  # Read back the actual value
        
        no_10 = xlookup(no_1, info_df.iloc[:, 9], info_df.iloc[:, 10], '')
        if col_no_10:
            set_if_empty(result_df, idx, col_no_10, no_10)
            no_10 = result_df.at[idx, col_no_10]  # Read back the actual value
        
        # ============================================================
        # FORMULA 11: İşveren Hakediş Birim Fiyat
        # ============================================================
        no_1_num = safe_float(no_1, 0)
        no_2_str = safe_str(no_2, '')
        
        if no_2_str in ['999-A', '999-C', '414-C'] or no_1_num == 313:
            isveren_hakedis_birim_fiyat = hourly_rate
        elif no_1_num in [312, 314, 316] or no_2_str == '360-T':
            isveren_hakedis_birim_fiyat = hourly_rate * 1.02
        elif no_2_str == '517-A':
            isveren_hakedis_birim_fiyat = safe_float(xlookup(person_id, info_df.iloc[:, 28], info_df.iloc[:, 33], 0))
        else:
            if summary_df is not None:
                val1 = safe_float(xlookup(no_1, summary_df.iloc[:, 2], summary_df.iloc[:, 26], 0))
                val2 = safe_float(xlookup(no_2, summary_df.iloc[:, 2], summary_df.iloc[:, 26], 0))
                isveren_hakedis_birim_fiyat = val1 + val2
            else:
                isveren_hakedis_birim_fiyat = 0
        
        if col_isveren_birim_fiyat:
            set_if_empty(result_df, idx, col_isveren_birim_fiyat, isveren_hakedis_birim_fiyat)
            isveren_hakedis_birim_fiyat = safe_float(result_df.at[idx, col_isveren_birim_fiyat])
        
        # ============================================================
        # FORMULA 12: İşveren-Hakediş
        # ============================================================
        if kuzey_mh_person > 0:
            isveren_hakedis = kuzey_mh_person * isveren_hakedis_birim_fiyat
        else:
            isveren_hakedis = isveren_hakedis_birim_fiyat * total_mh
        
        if col_isveren_hakedis:
            set_if_empty(result_df, idx, col_isveren_hakedis, isveren_hakedis)
            isveren_hakedis = safe_float(result_df.at[idx, col_isveren_hakedis])
        
        # ============================================================
        # FORMULA 13: İşveren Hakediş (USD)
        # ============================================================
        if isveren_currency == 'EURO':
            eur_usd_rate = safe_float(xlookup(week_month, info_df.iloc[:, 20], info_df.iloc[:, 23], 1))
            isveren_hakedis_usd = isveren_hakedis * eur_usd_rate
        else:
            isveren_hakedis_usd = isveren_hakedis
        
        if col_isveren_hakedis_usd:
            set_if_empty(result_df, idx, col_isveren_hakedis_usd, isveren_hakedis_usd)
            isveren_hakedis_usd = safe_float(result_df.at[idx, col_isveren_hakedis_usd])
        
        # ============================================================
        # FORMULA 14: İşveren Hakediş Birim Fiyatı (USD)
        # ============================================================
        if col_isveren_birim_fiyat_usd:
            if kuzey_mh_person > 0:
                isveren_hakedis_birim_fiyat_usd = isveren_hakedis_usd / kuzey_mh_person if kuzey_mh_person != 0 else 0
            else:
                isveren_hakedis_birim_fiyat_usd = isveren_hakedis_usd / total_mh if total_mh != 0 else 0
            set_if_empty(result_df, idx, col_isveren_birim_fiyat_usd, isveren_hakedis_birim_fiyat_usd)
        
        # ============================================================
        # FORMULA 15: Control-1
        # ============================================================
        if col_control_1:
            control_1 = xlookup(projects, info_df.iloc[:, 14], info_df.iloc[:, 18], '')
            set_if_empty(result_df, idx, col_control_1, control_1)
        
        # ============================================================
        # FORMULA 16: TM Liste
        # ============================================================
        if col_tm_liste:
            try:
                tm_liste = xlookup(person_id, info_df.iloc[:, 58], info_df.iloc[:, 60], '')
            except:
                tm_liste = ''
            set_if_empty(result_df, idx, col_tm_liste, tm_liste)
        
        # ============================================================
        # FORMULA 17: TM KOD
        # ============================================================
        if col_tm_kod:
            tm_kod = xlookup(projects, info_df.iloc[:, 14], info_df.iloc[:, 17], '')
            set_if_empty(result_df, idx, col_tm_kod, tm_kod)
        
        # ============================================================
        # FORMULA 18: Kontrol-1
        # ============================================================
        if col_kontrol_1:
            kontrol_1 = xlookup(projects, info_df.iloc[:, 47], info_df.iloc[:, 42], '')
            set_if_empty(result_df, idx, col_kontrol_1, kontrol_1)
        
        # ============================================================
        # FORMULA 19: Kontrol-2
        # ============================================================
        kontrol_2_str = "FALSE"

        if col_kontrol_2:
            kontrol_1_val = (
                result_df.at[idx, col_kontrol_1]
                if col_kontrol_1 is not None
                else ""

            )
            kontrol_2_str = "TRUE" if no_1 == kontrol_1_val else "FALSE"
            
            if idx == 0:
                print(f"\n=== DEBUG ROW 0 - Kontrol-2 ===")
                print(f"  col_kontrol_2 column name: '{col_kontrol_2}'")
                print(f"  no_1: {no_1}")
                print(f"  kontrol_1_val: {kontrol_1_val}")
                print(f"  kontrol_2_str: {kontrol_2_str}")
                print(f"  Current value at [{idx}, '{col_kontrol_2}']: {result_df.at[idx, col_kontrol_2]}")
                print("=" * 50)
            
            # Check if current value is numeric 1 or 0 (from previous incorrect fill)
            # If so, we should overwrite it with TRUE/FALSE
            current = result_df.at[idx, col_kontrol_2]
            should_overwrite = False
            
            if pd.isna(current):
                should_overwrite = True
            elif current in [0, 0.0, 1, 1.0]:
                # Overwrite numeric values with proper TRUE/FALSE
                should_overwrite = True
            elif isinstance(current, str) and current.strip() == '':
                should_overwrite = True
            
            if should_overwrite:
                # Convert column to object dtype if needed
                if result_df[col_kontrol_2].dtype in ['float64', 'int64']:
                    result_df[col_kontrol_2] = result_df[col_kontrol_2].astype('object')
                result_df.at[idx, col_kontrol_2] = kontrol_2_str
    
    print(f"Finished filling empty cells!")
    return result_df


@app.route('/api/process_empty_cells', methods=['POST'])
@login_required
def process_empty_cells():
    """
    Handle file upload with empty cells and fill them based on formulas.
    This uses the Info, Hourly Rates, and Summary sheets from previously uploaded files.
    """
    try:
        # Only admin can upload files
        if session.get('role') != 'admin':
            return jsonify({'error': 'Only admin can upload files'}), 403
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Check file extension
        allowed_extensions = {'.xlsx', '.xls', '.xlsb'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            return jsonify({'error': f'Invalid file type. Allowed: {", ".join(allowed_extensions)}'}), 400
        
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'{timestamp}_{filename}'
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Save file
        print(f"Saving file to: {filepath}")
        file.save(filepath)
        print(f"File saved successfully")
        
        # Load the DATABASE sheet from uploaded file
        print(f"Loading DATABASE sheet from uploaded file...")
        if file_ext == '.xlsb':
            df_database = pd.read_excel(filepath, sheet_name='DATABASE', engine='pyxlsb')
        else:
            df_database = pd.read_excel(filepath, sheet_name='DATABASE')
        
        print(f"DATABASE sheet loaded: {df_database.shape[0]} rows, {df_database.shape[1]} columns")
        
        # Load reference data (Info, Hourly Rates, Summary) from latest uploaded file
        print(f"Loading reference sheets (Info, Hourly Rates, Summary)...")
        if not load_excel_reference_data():
            return jsonify({'error': 'Could not load reference sheets (Info, Hourly Rates, Summary). Please ensure you have uploaded a file with these sheets first.'}), 400
        
        info_df = _excel_cache['info_df']
        rates_df = _excel_cache['hourly_rates_df']
        summary_df = _excel_cache['summary_df']
        
        print(f"Reference sheets loaded:")
        print(f"  - Info: {info_df.shape[0]} rows, {info_df.shape[1]} columns")
        print(f"  - Hourly Rates: {rates_df.shape[0]} rows, {rates_df.shape[1]} columns")
        print(f"  - Hourly Rates columns: {rates_df.columns.tolist()}")
        print(f"  - Hourly Rates column 0 (ID) first 5: {rates_df.iloc[:5, 0].tolist()}")
        print(f"  - Hourly Rates column 7 first 5: {rates_df.iloc[:5, 7].tolist()}")
        if summary_df is not None:
            print(f"  - Summary: {summary_df.shape[0]} rows, {summary_df.shape[1]} columns")
        else:
            print(f"  - Summary: Not available")
        
        # Fill empty cells based on formulas
        print(f"Filling empty cells based on formulas...")
        df_filled = fill_empty_cells_with_formulas(df_database, info_df, rates_df, summary_df)
        
        # Save the filled DataFrame back to Excel
        output_filename = f'filled_{timestamp}_{os.path.basename(file.filename)}'
        output_filepath = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        print(f"Saving filled data to: {output_filepath}")
        if file_ext == '.xlsb':
            # For xlsb, we need to save as xlsx since pyxlsb doesn't support writing
            output_filepath = output_filepath.replace('.xlsb', '.xlsx')
            df_filled.to_excel(output_filepath, sheet_name='DATABASE', index=False, engine='openpyxl')
        else:
            df_filled.to_excel(output_filepath, sheet_name='DATABASE', index=False)
        
        print(f"Filled file saved successfully!")
        
        return jsonify({
            'success': True,
            'message': f'File processed successfully! {df_filled.shape[0]} rows processed.',
            'original_file': filename,
            'filled_file': output_filename,
            'download_url': f'/api/download_filled/{output_filename}',
            'rows': df_filled.shape[0],
            'columns': df_filled.shape[1]
        })
    
    except Exception as e:
        print(f"Process empty cells error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/download_filled/<filename>')
@login_required
def download_filled(filename):
    """Download the filled Excel file"""
    try:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            return jsonify({'error': 'File not found'}), 404
        
        return send_file(filepath, as_attachment=True, download_name=filename)
    except Exception as e:
        print(f"Download error: {str(e)}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/filter', methods=['POST'])
@login_required
def filter_data():
    """Apply filters to data"""
    try:
        filters = request.json.get('filters', {})
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load combined data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                # Convert both to strings for consistent comparison
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        # Convert to JSON (replace NaN with empty string)
        df_clean = df.fillna('')
        data_json = df_clean.to_dict('records')
        
        return jsonify({
            'success': True,
            'data': data_json,
            'shape': df.shape
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/get-filtered-options', methods=['POST'])
@login_required
def get_filtered_options():
    """Get available filter options based on current filter selections (for cascading filters)"""
    try:
        filters = request.json.get('filters', {})
        
        print(f'DEBUG: Received filters: {filters}')
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load combined data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        print(f'DEBUG: Total rows before filtering: {len(df)}')
        
        # Apply filters progressively (cascading)
        for col, values in filters.items():
            if col in df.columns and values:
                # Convert both to strings for consistent comparison
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
                print(f'DEBUG: After filtering {col}: {len(df)} rows remain')
        
        print(f'DEBUG: Total rows after filtering: {len(df)}')
        
        # Get available options for each column after filtering
        # This creates the cascading effect - only show options that exist in filtered data
        filter_cols = []
        skip_cols = ['PERSONEL', 'id', 'created_at', 'updated_at']
        
        # Define preferred order for filters
        preferred_order = [
            'Name Surname',
            'Discipline',
            '(Week /\nMonth)',
            '(Week / Month)',
            'Week / Month',
            'Company',
            'Projects/Group',
            'Nationality',
            'Office Location',
            'Kuzey MH-Person',
            'Status',
            'North/\nSouth',
            'North/South',
            'Currency',
            'PP'
        ]
        
        # Process columns in preferred order first
        processed_cols = set()
        for pref_col in preferred_order:
            if pref_col in df.columns and pref_col not in skip_cols:
                try:
                    unique_count = df[pref_col].nunique()
                    
                    if unique_count > 0 and unique_count <= 500:
                        non_null_values = df[pref_col].dropna()
                        if len(non_null_values) > 0:
                            unique_values = sorted([str(v) for v in df[pref_col].dropna().unique()])
                            
                            filter_cols.append({
                                'name': pref_col,
                                'values': unique_values
                            })
                            processed_cols.add(pref_col)
                            print(f'DEBUG: {pref_col}: {len(unique_values)} unique values')
                except Exception as e:
                    print(f'DEBUG: Error processing {pref_col}: {e}')
                    continue
        
        # Add remaining columns
        for col in df.columns:
            if col in skip_cols or col in processed_cols:
                continue
                
            try:
                unique_count = df[col].nunique()
                
                if unique_count > 0 and unique_count <= 500:
                    non_null_values = df[col].dropna()
                    if len(non_null_values) > 0:
                        unique_values = sorted([str(v) for v in df[col].dropna().unique()])
                        
                        filter_cols.append({
                            'name': col,
                            'values': unique_values
                        })
                        print(f'DEBUG: {col}: {len(unique_values)} unique values')
            except Exception as e:
                print(f'DEBUG: Error processing {col}: {e}')
                continue
        
        print(f'DEBUG: Returning {len(filter_cols)} filter columns')
        
        return jsonify({
            'success': True,
            'filter_columns': filter_cols
        })
    
    except Exception as e:
        print(f'ERROR in get_filtered_options: {e}')
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/save-filter', methods=['POST'])
@login_required
def save_filter():
    """Save filter configuration for user"""
    try:
        data = request.json
        filter_name = data.get('filter_name')
        filter_type = data.get('filter_type')  # 'database', 'pivot', 'graph'
        filter_config = data.get('filter_config')
        
        if not filter_name or not filter_type or not filter_config:
            return jsonify({'error': 'Missing required fields'}), 400
        
        user_id = session.get('user_id')
        
        # Check if filter with same name and type exists for this user
        existing = SavedFilter.query.filter_by(
            user_id=user_id,
            filter_name=filter_name,
            filter_type=filter_type
        ).first()
        
        if existing:
            # Update existing filter
            existing.filter_config = json.dumps(filter_config)
            existing.updated_at = datetime.utcnow()
        else:
            # Create new filter
            new_filter = SavedFilter(
                user_id=user_id,
                filter_name=filter_name,
                filter_type=filter_type,
                filter_config=json.dumps(filter_config)
            )
            db.session.add(new_filter)
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Filter saved successfully'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/load-filters', methods=['GET'])
@login_required
def load_filters():
    """Load saved filters for user"""
    try:
        filter_type = request.args.get('filter_type')
        user_id = session.get('user_id')
        
        query = SavedFilter.query.filter_by(user_id=user_id)
        
        if filter_type:
            query = query.filter_by(filter_type=filter_type)
        
        filters = query.order_by(SavedFilter.updated_at.desc()).all()
        
        result = []
        for f in filters:
            result.append({
                'id': f.id,
                'filter_name': f.filter_name,
                'filter_type': f.filter_type,
                'filter_config': json.loads(f.filter_config),
                'created_at': f.created_at.isoformat(),
                'updated_at': f.updated_at.isoformat()
            })
        
        return jsonify({
            'success': True,
            'filters': result
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/delete-filter/<int:filter_id>', methods=['DELETE'])
@login_required
def delete_filter(filter_id):
    """Delete saved filter"""
    try:
        user_id = session.get('user_id')
        
        # Only allow user to delete their own filters
        saved_filter = SavedFilter.query.filter_by(
            id=filter_id,
            user_id=user_id
        ).first()
        
        if not saved_filter:
            return jsonify({'error': 'Filter not found'}), 404
        
        db.session.delete(saved_filter)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Filter deleted successfully'
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/pivot', methods=['POST'])
@login_required
def create_pivot():
    """Create pivot table"""
    try:
        config = request.json
        index_col = config.get('index')
        columns_col = config.get('columns')
        values_cols = config.get('values', [])
        agg_func = config.get('agg_func', 'sum')
        filters = config.get('filters', {})
        
        print(f"\n{'='*80}")
        print(f"PIVOT REQUEST DEBUG")
        print(f"{'='*80}")
        print(f"Index: {index_col}")
        print(f"Columns: {columns_col}")
        print(f"Values: {values_cols}")
        print(f"Agg Function: {agg_func}")
        print(f"Filters received: {len(filters)} categories")
        for col, vals in filters.items():
            print(f"  - {col}: {len(vals)} values")
        print(f"{'='*80}\n")
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        print(f"Data shape before filters: {df.shape}")
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                before_count = len(df)
                # Convert both dataframe column and filter values to strings for comparison
                # This ensures consistent comparison regardless of data type
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
                after_count = len(df)
                print(f"Filter '{col}': {before_count} → {after_count} rows (filtered out {before_count - after_count})")
                # Debug: Show what we're filtering
                if after_count == 0:
                    print(f"  WARNING: All data filtered out for column '{col}'")
                    print(f"  Filter values: {values_str[:10]}")
                    print(f"  Unique values in data: {df_col_str.unique()[:10].tolist()}")
        
        print(f"Data shape after filtering: {df.shape}")
        print(f"{'='*80}\n")
        # If no rows remain after filtering, return a clear error
        if df.empty:
            return jsonify({'error': 'No data available after applying filters. Please relax filters or select a different dataset.'}), 400
        
        # Create pivot
        if index_col and values_cols:
            # Validate that index column exists
            if index_col not in df.columns:
                return jsonify({'error': f'Index column "{index_col}" not found in data'}), 400
            
            # Validate and convert value columns to numeric
            valid_values = []
            for col in values_cols:
                print(f"\n--- Processing value column: '{col}' ---")
                
                if col not in df.columns:
                    print(f"❌ Column '{col}' not found in dataframe")
                    print(f"Available columns: {df.columns.tolist()}")
                    continue
                
                # Debug: Show sample values
                print(f"Column type: {df[col].dtype}")
                print(f"First 5 values: {df[col].head(5).tolist()}")
                print(f"Non-null count: {df[col].notna().sum()} / {len(df)}")
                    
                # Try to convert to numeric and ensure it's a Series
                try:
                    # Attempt robust numeric conversion: strip common currency/formatting
                    series = df[col]
                    # If object/string dtype, clean common characters like commas, currency symbols and words
                    if series.dtype == object or str(series.dtype).startswith('string'):
                        s_clean = series.astype(str)
                        # Remove common currency symbols and abbreviations, keep digits, dot and minus
                        s_clean = s_clean.str.replace(r'[\$,£€¥]', '', regex=True)
                        s_clean = s_clean.str.replace(r'\bUSD\b|\bTRY\b|\bEUR\b|\bTL\b', '', regex=True, case=False)
                        # Remove any non-numeric characters except dot and minus
                        s_clean = s_clean.str.replace(r'[^0-9.\-]', '', regex=True)
                        numeric_series = pd.to_numeric(s_clean, errors='coerce')
                    else:
                        numeric_series = pd.to_numeric(series, errors='coerce')
                    print(f"After cleaning + pd.to_numeric - type: {type(numeric_series)}, dtype: {numeric_series.dtype}")
                    
                    # Ensure it's 1-dimensional
                    if hasattr(numeric_series, 'ndim'):
                        print(f"Dimension check: ndim = {numeric_series.ndim}")
                        if numeric_series.ndim != 1:
                            print(f"❌ Column is {numeric_series.ndim}-dimensional, skipping")
                            continue
                    
                    # Check if we have any valid numeric values
                    valid_count = numeric_series.notna().sum()
                    null_count = numeric_series.isna().sum()
                    print(f"Valid numeric values: {valid_count}")
                    print(f"Null/NaN values: {null_count}")
                    print(f"Sample numeric values: {numeric_series.dropna().head(5).tolist()}")
                    
                    if valid_count == 0:
                        print(f"❌ No valid numeric values found")
                        continue
                    
                    # Update the dataframe with numeric values
                    df[col] = numeric_series
                    valid_values.append(col)
                    print(f"✅ Successfully converted '{col}' to numeric: {valid_count} valid values")
                except Exception as e:
                    print(f"❌ Error processing column '{col}': {str(e)}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            print(f"\n{'='*80}")
            print(f"VALID COLUMNS: {valid_values}")
            print(f"{'='*80}\n")
            
            if not valid_values:
                return jsonify({'error': 'No valid numeric columns selected for analysis. Please select columns with numeric values (costs, rates, etc.)'}), 400
            
            # Build pivot parameters
            pivot_params = {
                'values': valid_values[0] if len(valid_values) == 1 else valid_values,
                'index': index_col,
                'aggfunc': agg_func,
                'fill_value': 0  # Fill NaN with 0
            }
            
            # Add columns parameter if specified
            if columns_col and columns_col in df.columns:
                pivot_params['columns'] = columns_col
            
            print(f"Creating pivot with params: {pivot_params}")
            
            # Create pivot table
            pivot = pd.pivot_table(df, **pivot_params)
            pivot = pivot.reset_index()
            
            print(f"Pivot created successfully: {pivot.shape}")
            
            # Replace NaN with 0 for display
            pivot = pivot.fillna(0)
            
            # Format column names if they're tuples (multi-level columns)
            if any(isinstance(col, tuple) for col in pivot.columns):
                pivot.columns = [' - '.join(map(str, col)).strip(' -') if isinstance(col, tuple) else str(col) for col in pivot.columns]
            
            return jsonify({
                'success': True,
                'data': pivot.to_dict('records'),
                'columns': pivot.columns.tolist()
            })
        else:
            return jsonify({'error': 'Please select both Group By column and at least one Value column'}), 400
    
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Pivot error: {str(e)}")
        print(error_trace)
        return jsonify({'error': f'Error creating pivot table: {str(e)}'}), 500

@app.route('/api/chart', methods=['POST'])
@login_required
def create_chart():
    """Create chart from data"""
    try:
        config = request.json
        chart_type = config.get('chart_type')
        # Support both old and new parameter names
        x_col = config.get('x_column') or config.get('x')
        y_col = config.get('y_column') or config.get('y')
        color_col = config.get('color_column') or config.get('color')
        filters = config.get('filters', {})
        
        # Always use whole database for graphs (no user filter)
        user_filter = None
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        
        if df.empty:
            return jsonify({'error': 'No data available'}), 400
        
        # Apply filters BEFORE adding calculated columns for better performance
        if filters:
            for col, values in filters.items():
                if col in df.columns and values:
                    df = df[df[col].isin(values)]
        
        # Add calculated columns to ALL data (don't limit yet)
        df = add_calculated_columns(df)
        
        # Validate columns exist
        if x_col and x_col not in df.columns:
            return jsonify({'error': f'Column "{x_col}" not found'}), 400
        if y_col and y_col not in df.columns:
            return jsonify({'error': f'Column "{y_col}" not found'}), 400
        
        # Clean data for plotting - handle mixed types
        if x_col:
            df[x_col] = df[x_col].fillna('')
            # No limit on unique values - show all data
        
        if y_col:
            # Convert y column to numeric if possible
            df[y_col] = pd.to_numeric(df[y_col], errors='coerce').fillna(0)
            # Remove rows with zero or negative values for pie charts
            if chart_type == 'pie':
                df = df[df[y_col] > 0]
        
        if color_col and color_col in df.columns:
            df[color_col] = df[color_col].fillna('')
        
        # Remove rows with invalid data
        df = df.dropna(subset=[col for col in [x_col, y_col] if col and col in df.columns])
        
        if df.empty:
            return jsonify({'error': 'No valid data after filtering'}), 400
        
        # For line charts: prefer date format over week codes if both exist
        if chart_type == 'line' and x_col and x_col in df.columns:
            import re
            # Check if we have date format entries (containing / and NOT matching W##)
            df[x_col] = df[x_col].astype(str)
            has_dates = df[x_col].str.contains('/', na=False).any()
            has_week_codes = df[x_col].str.match(r'^W\d+$', case=False, na=False).any()
            
            # Only filter out week codes if we also have date format entries
            if has_dates and has_week_codes:
                # Keep only rows that contain "/" (date format)
                df = df[df[x_col].str.contains('/', na=False)]
        
        if df.empty:
            return jsonify({'error': 'No valid data after filtering'}), 400
        
        # Create chart
        fig = None
        color_param = color_col if color_col and color_col in df.columns else None
        
        # Sort data by X column if it looks like a date (for proper ordering)
        x_sort_col = None
        if x_col:
            try:
                # Create a temporary datetime column for sorting only
                x_sort_col = f'{x_col}_sort_temp'
                import warnings
                with warnings.catch_warnings():
                    warnings.filterwarnings('ignore', category=UserWarning)
                    df[x_sort_col] = pd.to_datetime(df[x_col], errors='coerce', format='mixed')
                if df[x_sort_col].notna().any():
                    df = df.sort_values(x_sort_col)
                    df = df.drop(columns=[x_sort_col])
            except:
                pass
        
        # Use simpler chart types for better performance
        if chart_type == 'bar':
            # Aggregate data by grouping - always use SUM for totals
            if color_param:
                df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
                fig = px.bar(df_agg, x=x_col, y=y_col, color=color_param)
            else:
                df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
                fig = px.bar(df_agg, x=x_col, y=y_col)
        elif chart_type == 'line':
            # Aggregate data properly by grouping X and Color columns
            if color_param:
                # Group by both X axis and Color, then sum the Y values
                df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
                # Ensure data types are correct
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').fillna(0)
                
                # Sort by date if x_col looks like dates (contains /)
                if df_agg[x_col].astype(str).str.contains('/', na=False).any():
                    try:
                        # Try multiple date formats: dd/mmm, dd/mm/yyyy, etc.
                        # Add year 2000 for dates without year (01/Jan becomes 01/Jan/2000)
                        def parse_date(date_str):
                            if pd.isna(date_str):
                                return pd.NaT
                            date_str = str(date_str).strip()
                            # Try with year first
                            for fmt in ['%d/%m/%Y', '%d/%b/%Y', '%d-%m-%Y', '%d-%b-%Y']:
                                try:
                                    return pd.to_datetime(date_str, format=fmt)
                                except:
                                    pass
                            # Try without year (add 2000 as default year)
                            for fmt in ['%d/%b', '%d/%m']:
                                try:
                                    parsed = pd.to_datetime(date_str + '/2000', format=fmt + '/%Y')
                                    return parsed
                                except:
                                    pass
                            return pd.NaT
                        
                        df_agg['_temp_date'] = df_agg[x_col].apply(parse_date)
                        if df_agg['_temp_date'].notna().any():
                            df_agg = df_agg.sort_values(['_temp_date', color_param])
                            df_agg = df_agg.drop(columns=['_temp_date'])
                        else:
                            df_agg = df_agg.sort_values([color_param, x_col])
                    except Exception as e:
                        print(f"Date parsing error: {e}")
                        df_agg = df_agg.sort_values([color_param, x_col])
                else:
                    df_agg = df_agg.sort_values([color_param, x_col])
                    
                fig = px.line(df_agg, x=x_col, y=y_col, color=color_param, markers=True)
            else:
                # Group by X axis only, then sum the Y values
                df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
                # Ensure data types are correct
                df_agg[y_col] = pd.to_numeric(df_agg[y_col], errors='coerce').fillna(0)
                
                # Sort by date if x_col looks like dates
                if df_agg[x_col].astype(str).str.contains('/', na=False).any():
                    try:
                        # Try multiple date formats
                        def parse_date(date_str):
                            if pd.isna(date_str):
                                return pd.NaT
                            date_str = str(date_str).strip()
                            # Try with year first
                            for fmt in ['%d/%m/%Y', '%d/%b/%Y', '%d-%m-%Y', '%d-%b-%Y']:
                                try:
                                    return pd.to_datetime(date_str, format=fmt)
                                except:
                                    pass
                            # Try without year (add 2000 as default)
                            for fmt in ['%d/%b', '%d/%m']:
                                try:
                                    parsed = pd.to_datetime(date_str + '/2000', format=fmt + '/%Y')
                                    return parsed
                                except:
                                    pass
                            return pd.NaT
                        
                        df_agg['_temp_date'] = df_agg[x_col].apply(parse_date)
                        if df_agg['_temp_date'].notna().any():
                            df_agg = df_agg.sort_values('_temp_date')
                            df_agg = df_agg.drop(columns=['_temp_date'])
                        else:
                            df_agg = df_agg.sort_values(x_col)
                    except Exception as e:
                        print(f"Date parsing error: {e}")
                        df_agg = df_agg.sort_values(x_col)
                else:
                    df_agg = df_agg.sort_values(x_col)
                    
                fig = px.line(df_agg, x=x_col, y=y_col, markers=True)
        elif chart_type == 'scatter':
            # Scatter plots show individual points, but can still aggregate
            if color_param:
                df_agg = df.groupby([x_col, color_param], as_index=False)[y_col].sum()
                fig = px.scatter(df_agg, x=x_col, y=y_col, color=color_param)
            else:
                df_agg = df.groupby(x_col, as_index=False)[y_col].sum()
                fig = px.scatter(df_agg, x=x_col, y=y_col)
        elif chart_type == 'pie':
            # Aggregate for pie chart
            df_pie = df.groupby(x_col)[y_col].sum().reset_index()
            # Limit to top 10 slices
            df_pie = df_pie.nlargest(10, y_col)
            fig = px.pie(df_pie, names=x_col, values=y_col)
        elif chart_type == 'box':
            fig = px.box(df, x=x_col, y=y_col, color=color_param)
        elif chart_type == 'histogram':
            fig = px.histogram(df, x=x_col, y=y_col, color=color_param)
        else:
            return jsonify({'error': f'Unsupported chart type: {chart_type}'}), 400
        
        if fig:
            # Update layout for better appearance and performance
            fig.update_layout(
                template='plotly_white',
                title=f'{chart_type.upper()} Chart',
                xaxis_title=x_col,
                yaxis_title=y_col,
                height=500,
                showlegend=True if color_param else False
            )
            
            # Convert to JSON string - use Plotly's built-in method which handles numpy arrays
            try:
                import plotly.utils
                chart_json = plotly.utils.PlotlyJSONEncoder().encode(fig)
            except Exception as encode_err:
                print(f"Plotly encoding error: {str(encode_err)}")
                # Fallback: try to_json method
                try:
                    chart_json = fig.to_json()
                except Exception as json_err:
                    print(f"Plotly to_json error: {str(json_err)}")
                    return jsonify({'error': f'Failed to encode chart: {str(encode_err)}'}), 500
            
            return jsonify({
                'success': True,
                'chart': chart_json
            })
        else:
            return jsonify({'error': 'Failed to create chart'}), 500
    
    except Exception as e:
        print(f"Chart error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/favorites', methods=['GET'])
def get_favorites():
    """Get list of favorite reports"""
    data = load_favorites()
    return jsonify(data)

@app.route('/api/favorites', methods=['POST'])
def save_favorite():
    """Save a favorite report"""
    try:
        report_name = request.json.get('name')
        report_config = request.json.get('config')
        
        data = load_favorites()
        data['reports'][report_name] = report_config
        save_favorites(data['reports'], report_name)
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/favorites/<name>', methods=['DELETE'])
def delete_favorite(name):
    """Delete a favorite report"""
    try:
        data = load_favorites()
        if name in data['reports']:
            del data['reports'][name]
            last_loaded = data['_last_loaded'] if data['_last_loaded'] != name else None
            save_favorites(data['reports'], last_loaded)
            return jsonify({'success': True})
        else:
            return jsonify({'error': 'Report not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export', methods=['POST'])
@login_required
def export_report():
    """Export comprehensive report to Word or Excel with data, charts, and pivot tables"""
    try:
        export_format = request.json.get('format', 'excel')
        filters = request.json.get('filters', {})
        pivot_config = request.json.get('pivot_config', None)
        chart_configs = request.json.get('chart_configs', [])
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data from database
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        if df.empty:
            return jsonify({'error': 'No data available to export'}), 400
        
        if export_format == 'word':
            # Create comprehensive Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('📊 Data Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date and metadata
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            date_run.font.size = Pt(11)
            
            user_para = doc.add_paragraph()
            user_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            user_run = user_para.add_run(f'Generated by: {session.get("name", "User")}')
            user_run.font.size = Pt(10)
            
            doc.add_page_break()
            
            # Dataset Overview
            doc.add_heading('1. Dataset Overview', 1)
            doc.add_paragraph(f'📋 Total Rows: {df.shape[0]:,}')
            doc.add_paragraph(f'📊 Total Columns: {df.shape[1]}')
            
            # Unique staff count
            if 'Name Surname' in df.columns:
                unique_staff = df['Name Surname'].nunique()
                doc.add_paragraph(f'👥 Unique Staff Members: {unique_staff}')
            
            # Active filters
            if filters:
                doc.add_heading('Active Filters:', 2)
                for col, values in filters.items():
                    doc.add_paragraph(f'• {col}: {len(values)} selection(s)', style='List Bullet')
            
            doc.add_page_break()
            
            # Sample data table (first 20 rows)
            doc.add_heading('2. Sample Data (First 20 Rows)', 1)
            
            # Limit columns for readability in Word
            display_cols = df.columns[:10].tolist()  # First 10 columns
            sample_df = df[display_cols].head(20)
            
            # Create table
            table = doc.add_table(rows=len(sample_df) + 1, cols=len(display_cols))
            table.style = 'Light Grid Accent 1'
            
            # Header row
            for i, col in enumerate(display_cols):
                table.rows[0].cells[i].text = str(col)
                table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for row_idx, (_, row) in enumerate(sample_df.iterrows(), start=1):
                for col_idx, col in enumerate(display_cols):
                    cell_value = str(row[col]) if pd.notna(row[col]) else ''
                    table.rows[row_idx].cells[col_idx].text = cell_value
            
            # Pivot table with actual data
            if pivot_config:
                doc.add_page_break()
                doc.add_heading('3. Pivot Analysis', 1)
                
                # Configuration info
                doc.add_paragraph(f'📊 Index (Group By): {pivot_config.get("index", "N/A")}')
                doc.add_paragraph(f'📋 Columns: {pivot_config.get("columns", "None")}')
                doc.add_paragraph(f'📈 Values: {", ".join(pivot_config.get("values", []))}')
                doc.add_paragraph(f'🔢 Aggregation: {pivot_config.get("agg_func", "sum")}')
                doc.add_paragraph('')
                
                # Create actual pivot table
                try:
                    index_col = pivot_config.get('index')
                    columns_col = pivot_config.get('columns')
                    values_cols = pivot_config.get('values', [])
                    agg_func = pivot_config.get('agg_func', 'sum')
                    
                    if index_col and values_cols:
                        # Convert numeric columns
                        for col in values_cols:
                            if col in df.columns:
                                df[col] = pd.to_numeric(df[col], errors='coerce')
                        
                        pivot_df = pd.pivot_table(
                            df,
                            values=values_cols[0] if len(values_cols) == 1 else values_cols,
                            index=index_col,
                            columns=columns_col if columns_col else None,
                            aggfunc=agg_func,
                            fill_value=0
                        )
                        
                        # Reset index to make it a regular column
                        pivot_df = pivot_df.reset_index()
                        
                        # Limit to first 30 rows for readability
                        display_pivot = pivot_df.head(30)
                        
                        # Create Word table
                        pivot_table = doc.add_table(rows=len(display_pivot) + 1, cols=len(display_pivot.columns))
                        pivot_table.style = 'Light Grid Accent 1'
                        
                        # Header row
                        for i, col in enumerate(display_pivot.columns):
                            cell = pivot_table.rows[0].cells[i]
                            cell.text = str(col)
                            cell.paragraphs[0].runs[0].font.bold = True
                        
                        # Data rows
                        for row_idx, (_, row) in enumerate(display_pivot.iterrows(), start=1):
                            for col_idx, col in enumerate(display_pivot.columns):
                                value = row[col]
                                if isinstance(value, (int, float)):
                                    cell_text = f'{value:,.2f}' if not pd.isna(value) else '0'
                                else:
                                    cell_text = str(value) if pd.notna(value) else ''
                                pivot_table.rows[row_idx].cells[col_idx].text = cell_text
                        
                        if len(pivot_df) > 30:
                            doc.add_paragraph(f'\n(Showing first 30 of {len(pivot_df)} total rows. Full table in Excel export.)')
                except Exception as e:
                    doc.add_paragraph(f'Error creating pivot table: {str(e)}')
                    doc.add_paragraph('Full pivot table available in Excel export.')
            
            # Chart configurations and data
            if chart_configs:
                doc.add_page_break()
                doc.add_heading('4. Chart Data', 1)
                doc.add_paragraph(f'Total Charts Created: {len(chart_configs)}')
                doc.add_paragraph('')
                
                for i, config in enumerate(chart_configs, 1):
                    chart_type = config.get("chart_type", "Unknown").title()
                    x_col = config.get("x_column", "")
                    y_col = config.get("y_column", "")
                    color_col = config.get("color_column", "")
                    
                    doc.add_heading(f'Chart {i}: {chart_type}', 2)
                    
                    # Configuration section
                    if chart_type.lower() == 'pie':
                        doc.add_paragraph(f'📊 Labels (Categories): {x_col}')
                        doc.add_paragraph(f'📈 Values (Sizes): {y_col}')
                    else:
                        doc.add_paragraph(f'📊 X-axis: {x_col}')
                        doc.add_paragraph(f'📈 Y-axis: {y_col}')
                    
                    if color_col:
                        doc.add_paragraph(f'🎨 Color by: {color_col}')
                    doc.add_paragraph('')
                    
                    # Generate chart data table
                    try:
                        # Prepare data for chart based on type
                        if x_col in df.columns and y_col in df.columns:
                            # Select columns based on what's configured
                            cols_to_use = [x_col, y_col]
                            
                            # For bar/pie charts, we need to aggregate by category
                            if chart_type.lower() in ['bar', 'pie']:
                                chart_df = df[[x_col, y_col]].copy()
                                chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                                # Aggregate by X column (category)
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                                chart_data = chart_data.sort_values(by=f'Total {y_col}', ascending=False).head(25)
                            
                            # For line charts, show time series data
                            elif chart_type.lower() == 'line':
                                chart_df = df[[x_col, y_col]].copy()
                                chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                                # Group by X column (time) and aggregate
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                                # Try to sort by date if possible
                                try:
                                    chart_data[x_col] = pd.to_datetime(chart_data[x_col], errors='coerce')
                                    chart_data = chart_data.sort_values(by=x_col)
                                    chart_data[x_col] = chart_data[x_col].dt.strftime('%d/%b/%Y')
                                except:
                                    pass
                                chart_data = chart_data.head(25)
                            
                            # For scatter and other types, show raw data points
                            else:
                                chart_df = df[[x_col, y_col]].copy()
                                chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                                chart_data = chart_df.dropna().head(50)
                            
                            # Create Word table for chart data
                            if not chart_data.empty:
                                doc.add_paragraph('Data Summary:', style='Heading 3')
                                
                                # Limit rows for display
                                display_chart = chart_data.head(25)
                                chart_table = doc.add_table(rows=len(display_chart) + 1, cols=len(display_chart.columns))
                                chart_table.style = 'Light Grid Accent 1'
                                
                                # Header row
                                for col_idx, col in enumerate(display_chart.columns):
                                    cell = chart_table.rows[0].cells[col_idx]
                                    cell.text = str(col)
                                    cell.paragraphs[0].runs[0].font.bold = True
                                
                                # Data rows
                                for row_idx, (_, row) in enumerate(display_chart.iterrows(), start=1):
                                    for col_idx, col in enumerate(display_chart.columns):
                                        value = row[col]
                                        if isinstance(value, (int, float)):
                                            cell_text = f'{value:,.2f}' if not pd.isna(value) else '0'
                                        else:
                                            cell_text = str(value) if pd.notna(value) else ''
                                        chart_table.rows[row_idx].cells[col_idx].text = cell_text
                                
                                if len(chart_data) > 25:
                                    doc.add_paragraph(f'(Showing top 25 of {len(chart_data)} data points)')
                                
                                doc.add_paragraph('')
                            else:
                                doc.add_paragraph('⚠️ No data available for this chart.')
                        else:
                            doc.add_paragraph(f'⚠️ Missing columns: {x_col} or {y_col}')
                    except Exception as e:
                        doc.add_paragraph(f'⚠️ Error generating chart data: {str(e)}')
                    
                    doc.add_paragraph('')
                
                doc.add_paragraph('Note: Interactive chart visualizations are available in the web application. ')
                doc.add_paragraph('This report includes the underlying chart data for analysis.')
            
            # Save Word document
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return send_file(
                doc_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        
        else:  # Excel export with multiple sheets
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Sheet 1: Full filtered data
                df.to_excel(writer, sheet_name='Filtered Data', index=False)
                
                # Format the data sheet
                worksheet = writer.sheets['Filtered Data']
                for i, col in enumerate(df.columns):
                    max_len = max(df[col].astype(str).apply(len).max(), len(str(col))) + 2
                    worksheet.set_column(i, i, min(max_len, 50))
                
                # Sheet 2: Summary statistics
                summary_df = df.describe(include='all').transpose()
                summary_df.to_excel(writer, sheet_name='Summary Statistics')
                
                # Sheet 3: Pivot table (if config provided)
                if pivot_config:
                    try:
                        index_col = pivot_config.get('index')
                        columns_col = pivot_config.get('columns')
                        values_cols = pivot_config.get('values', [])
                        agg_func = pivot_config.get('agg_func', 'sum')
                        
                        if index_col and values_cols:
                            pivot_df = pd.pivot_table(
                                df,
                                values=values_cols[0] if len(values_cols) == 1 else values_cols,
                                index=index_col,
                                columns=columns_col if columns_col else None,
                                aggfunc=agg_func,
                                fill_value=0
                            )
                            pivot_df.to_excel(writer, sheet_name='Pivot Table')
                    except Exception as e:
                        print(f"Error creating pivot in Excel: {e}")
                
                # Sheet 4: Metadata
                metadata = {
                    'Export Date': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    'Generated By': [session.get('name', 'User')],
                    'Total Rows': [df.shape[0]],
                    'Total Columns': [df.shape[1]],
                    'Filters Applied': [len(filters)]
                }
                metadata_df = pd.DataFrame(metadata)
                metadata_df.to_excel(writer, sheet_name='Metadata', index=False)
                
                # Sheet 5: Chart configurations (if provided)
                if chart_configs:
                    chart_data = []
                    for i, config in enumerate(chart_configs, 1):
                        chart_data.append({
                            'Chart Number': i,
                            'Type': config.get('chart_type', ''),
                            'X-axis': config.get('x_column', ''),
                            'Y-axis': config.get('y_column', ''),
                            'Color': config.get('color_column', 'None')
                        })
                    charts_df = pd.DataFrame(chart_data)
                    charts_df.to_excel(writer, sheet_name='Chart Configs', index=False)
            
            excel_io.seek(0)
            
            return send_file(
                excel_io,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
    
    except Exception as e:
        print(f"Export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/export-pivot', methods=['POST'])
@login_required
def export_pivot():
    """Export pivot table to Excel or Word"""
    try:
        print("=== PIVOT EXPORT DEBUG ===")
        export_format = request.json.get('format', 'excel')
        filters = request.json.get('filters', {})
        pivot_config = request.json.get('pivot_config', None)
        
        print(f"Export format: {export_format}")
        print(f"Pivot config: {pivot_config}")
        
        if not pivot_config:
            return jsonify({'error': 'No pivot configuration provided'}), 400
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        if df.empty:
            return jsonify({'error': 'No data available to export'}), 400
        
        # Create pivot table
        index_col = pivot_config.get('index')
        columns_col = pivot_config.get('columns')
        values_cols = pivot_config.get('values')
        agg_func = pivot_config.get('agg_func', 'sum')
        
        # Handle both single value and list of values
        if isinstance(values_cols, str):
            values_cols = [values_cols]
        elif not isinstance(values_cols, list):
            values_cols = [str(values_cols)] if values_cols else []
        
        print(f"Index: {index_col}, Columns: {columns_col}, Values: {values_cols}, Agg: {agg_func}")
        
        if not index_col or not values_cols:
            return jsonify({'error': 'Invalid pivot configuration - missing index or values'}), 400
        
        # Build pivot
        try:
            # Validate and convert all value columns to numeric
            valid_values = []
            for val_col in values_cols:
                if val_col in df.columns:
                    df[val_col] = pd.to_numeric(df[val_col], errors='coerce')
                    valid_values.append(val_col)
                    print(f"Converted {val_col} to numeric, non-null count: {df[val_col].notna().sum()}")
                else:
                    print(f"Warning: Value column '{val_col}' not found in data")
            
            if not valid_values:
                return jsonify({'error': 'None of the selected value columns were found in the data'}), 400
            
            print(f"Using value columns: {valid_values}")
            print(f"Available columns: {df.columns.tolist()[:10]}")
            
            # Use all valid value columns (or just first if only one)
            values_to_use = valid_values[0] if len(valid_values) == 1 else valid_values
            
            pivot_params = {
                'values': values_to_use,
                'index': index_col,
                'aggfunc': agg_func,
                'fill_value': 0
            }
            
            if columns_col and columns_col in df.columns:
                pivot_params['columns'] = columns_col
                print(f"Added columns parameter: {columns_col}")
            
            pivot_df = pd.pivot_table(df, **pivot_params)
            pivot_df = pivot_df.reset_index()
            
            # Handle multi-level columns
            if any(isinstance(col, tuple) for col in pivot_df.columns):
                pivot_df.columns = [' - '.join(map(str, col)).strip(' -') if isinstance(col, tuple) else str(col) for col in pivot_df.columns]
            
        except Exception as e:
            return jsonify({'error': f'Error creating pivot: {str(e)}'}), 500
        
        if export_format == 'word':
            # Create Word document with pivot table
            doc = Document()
            
            # Title
            title = doc.add_heading('📊 Pivot Table Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            date_run.font.size = Pt(11)
            
            doc.add_paragraph()
            
            # Configuration
            doc.add_heading('Pivot Configuration', 1)
            doc.add_paragraph(f'📊 Rows (Index): {index_col}')
            if columns_col:
                doc.add_paragraph(f'📋 Columns: {columns_col}')
            doc.add_paragraph(f'📈 Values: {", ".join(valid_values)}')
            doc.add_paragraph(f'🔢 Aggregation: {agg_func.upper()}')
            
            doc.add_paragraph()
            
            # Pivot table data
            doc.add_heading('Pivot Table', 1)
            
            # Limit to reasonable size for Word
            display_pivot = pivot_df.head(100)
            
            # Create table
            pivot_table = doc.add_table(rows=len(display_pivot) + 1, cols=len(display_pivot.columns))
            pivot_table.style = 'Light Grid Accent 1'
            
            # Header row
            for i, col in enumerate(display_pivot.columns):
                cell = pivot_table.rows[0].cells[i]
                cell.text = str(col)
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for row_idx, (_, row) in enumerate(display_pivot.iterrows(), start=1):
                for col_idx, col in enumerate(display_pivot.columns):
                    value = row[col]
                    if isinstance(value, (int, float)):
                        cell_text = f'{value:,.2f}' if not pd.isna(value) else '0'
                    else:
                        cell_text = str(value) if pd.notna(value) else ''
                    pivot_table.rows[row_idx].cells[col_idx].text = cell_text
            
            if len(pivot_df) > 100:
                doc.add_paragraph(f'\n(Showing first 100 of {len(pivot_df)} total rows)')
            
            # Save Word document
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return send_file(
                doc_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'pivot_table_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        
        else:  # Excel export
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Sheet 1: Pivot table
                pivot_df.to_excel(writer, sheet_name='Pivot Table', index=False)
                
                # Format the pivot sheet
                worksheet = writer.sheets['Pivot Table']
                for i, col in enumerate(pivot_df.columns):
                    max_len = max(pivot_df[col].astype(str).apply(len).max(), len(str(col))) + 2
                    worksheet.set_column(i, i, min(max_len, 50))
                
                # Sheet 2: Filtered data (source)
                df.to_excel(writer, sheet_name='Source Data', index=False)
                
                # Sheet 3: Configuration
                config_df = pd.DataFrame({
                    'Setting': ['Index/Rows', 'Columns', 'Values', 'Aggregation'],
                    'Value': [index_col, columns_col or 'None', ', '.join(valid_values), agg_func.upper()]
                })
                config_df.to_excel(writer, sheet_name='Configuration', index=False)
            
            excel_io.seek(0)
            
            return send_file(
                excel_io,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'pivot_table_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
    
    except Exception as e:
        print(f"Pivot export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/export-charts', methods=['POST'])
@login_required
def export_charts():
    """Export chart data to Excel or Word"""
    try:
        export_format = request.json.get('format', 'excel')
        filters = request.json.get('filters', {})
        chart_configs = request.json.get('chart_configs', [])
        
        if not chart_configs:
            return jsonify({'error': 'No chart configurations provided'}), 400
        
        # Get user filter
        user_filter = None if session.get('role') == 'admin' else session.get('name')
        
        # Load and filter data
        file_path = session.get('current_file')
        df = get_combined_data(file_path, user_filter)
        df = add_calculated_columns(df)
        
        # Apply filters
        for col, values in filters.items():
            if col in df.columns and values:
                df_col_str = df[col].astype(str)
                values_str = [str(v) for v in values]
                df = df[df_col_str.isin(values_str)]
        
        if df.empty:
            return jsonify({'error': 'No data available to export'}), 400
        
        if export_format == 'word':
            # Create Word document with chart images
            doc = Document()
            
            # Title
            title = doc.add_heading('📈 Charts Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
            date_run.font.size = Pt(11)
            
            doc.add_paragraph()
            doc.add_heading(f'Total Charts: {len(chart_configs)}', 1)
            doc.add_paragraph()
            
            # Process each chart
            for i, config in enumerate(chart_configs, 1):
                chart_type = config.get("chart_type", "Unknown").title()
                x_col = config.get("x_column", "")
                y_col = config.get("y_column", "")
                color_col = config.get("color_column", "")
                
                doc.add_heading(f'Chart {i}: {chart_type}', 2)
                
                # Configuration
                if chart_type.lower() == 'pie':
                    doc.add_paragraph(f'📊 Categories: {x_col}')
                    doc.add_paragraph(f'📈 Values: {y_col}')
                else:
                    doc.add_paragraph(f'📊 X-axis: {x_col}')
                    doc.add_paragraph(f'📈 Y-axis: {y_col}')
                
                if color_col:
                    doc.add_paragraph(f'🎨 Color by: {color_col}')
                
                doc.add_paragraph()
                
                # Generate chart image or data table
                image_added = False
                chart_data = None
                
                try:
                    if x_col in df.columns and y_col in df.columns:
                        # Prepare chart data
                        chart_df = df[[x_col, y_col]].copy()
                        chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                        chart_df = chart_df.dropna()
                        
                        # Aggregate based on chart type
                        if chart_type.lower() in ['bar', 'pie']:
                            chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                            chart_data = chart_data.sort_values(by=y_col, ascending=False).head(30)
                        elif chart_type.lower() == 'line':
                            chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                            # Try to sort by date
                            try:
                                chart_data[x_col] = pd.to_datetime(chart_data[x_col], errors='coerce')
                                chart_data = chart_data.sort_values(by=x_col)
                                # Convert back to string for display
                                chart_data[x_col] = chart_data[x_col].dt.strftime('%d/%b/%Y')
                            except:
                                pass
                            chart_data = chart_data.head(50)
                        else:
                            chart_data = chart_df.head(50)
                        
                        # Try to create chart image
                        try:
                            # Create Plotly figure
                            fig = None
                            if chart_type.lower() == 'bar':
                                fig = px.bar(chart_data, x=x_col, y=y_col, title=f'{chart_type} Chart: {y_col} by {x_col}')
                            elif chart_type.lower() == 'line':
                                fig = px.line(chart_data, x=x_col, y=y_col, title=f'{chart_type} Chart: {y_col} over {x_col}')
                            elif chart_type.lower() == 'pie':
                                fig = px.pie(chart_data, names=x_col, values=y_col, title=f'{chart_type} Chart: {y_col} by {x_col}')
                            elif chart_type.lower() == 'scatter':
                                fig = px.scatter(chart_data, x=x_col, y=y_col, title=f'{chart_type} Chart: {y_col} vs {x_col}')
                            
                            if fig:
                                # Update layout for better export
                                fig.update_layout(
                                    width=800,
                                    height=500,
                                    template='plotly_white',
                                    showlegend=True
                                )
                                
                                # Export to image using kaleido
                                img_bytes = fig.to_image(format='png', width=800, height=500, scale=2)
                                img_stream = io.BytesIO(img_bytes)
                                doc.add_picture(img_stream, width=Inches(6))
                                image_added = True
                                doc.add_paragraph()
                        except Exception as img_error:
                            print(f"Chart image generation failed: {str(img_error)}")
                            # Will fall back to table below
                        
                        # If image failed, show data table as fallback
                        if not image_added and chart_data is not None and not chart_data.empty:
                            doc.add_paragraph('Chart Data Table:', style='Heading 3')
                            
                            # Create table
                            table = doc.add_table(rows=min(len(chart_data) + 1, 31), cols=2)
                            table.style = 'Light Grid Accent 1'
                            
                            # Header
                            table.rows[0].cells[0].text = str(x_col)
                            table.rows[0].cells[1].text = str(y_col)
                            for cell in table.rows[0].cells:
                                cell.paragraphs[0].runs[0].font.bold = True
                            
                            # Data rows (max 30)
                            for idx, (_, row) in enumerate(chart_data.head(30).iterrows(), start=1):
                                table.rows[idx].cells[0].text = str(row[x_col]) if pd.notna(row[x_col]) else ''
                                value = row[y_col]
                                table.rows[idx].cells[1].text = f'{value:,.2f}' if pd.notna(value) else '0'
                            
                            doc.add_paragraph()
                    else:
                        doc.add_paragraph(f'⚠️ Missing columns: {x_col} or {y_col}')
                
                except Exception as e:
                    doc.add_paragraph(f'⚠️ Error generating chart: {str(e)}')
                
                doc.add_paragraph()
            
            # Save Word document
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            return send_file(
                doc_io,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'charts_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            )
        
        else:  # Excel export
            excel_io = io.BytesIO()
            with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Sheet 1: Chart configurations
                chart_config_data = []
                for i, config in enumerate(chart_configs, 1):
                    chart_config_data.append({
                        'Chart #': i,
                        'Type': config.get('chart_type', ''),
                        'X-axis': config.get('x_column', ''),
                        'Y-axis': config.get('y_column', ''),
                        'Color': config.get('color_column', 'None')
                    })
                config_df = pd.DataFrame(chart_config_data)
                config_df.to_excel(writer, sheet_name='Chart Configurations', index=False)
                
                # Sheets 2+: Data for each chart
                for i, config in enumerate(chart_configs, 1):
                    try:
                        x_col = config.get("x_column", "")
                        y_col = config.get("y_column", "")
                        chart_type = config.get("chart_type", "").lower()
                        
                        if x_col in df.columns and y_col in df.columns:
                            chart_df = df[[x_col, y_col]].copy()
                            chart_df[y_col] = pd.to_numeric(chart_df[y_col], errors='coerce')
                            
                            # Aggregate based on type
                            if chart_type in ['bar', 'pie']:
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                                chart_data = chart_data.sort_values(by=f'Total {y_col}', ascending=False)
                            elif chart_type == 'line':
                                chart_data = chart_df.groupby(x_col)[y_col].sum().reset_index()
                                chart_data.columns = [x_col, f'Total {y_col}']
                            else:
                                chart_data = chart_df.dropna().head(100)
                            
                            # Save to sheet
                            sheet_name = f'Chart {i} Data'[:31]  # Excel sheet name limit
                            chart_data.to_excel(writer, sheet_name=sheet_name, index=False)
                    except:
                        continue
                
                # Last sheet: Filtered source data
                df.to_excel(writer, sheet_name='Source Data', index=False)
            
            excel_io.seek(0)
            
            return send_file(
                excel_io,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f'charts_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
            )
    
    except Exception as e:
        print(f"Charts export error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    with app.app_context():
        init_db()
    app.run(debug=True, host='0.0.0.0', port=5000)




