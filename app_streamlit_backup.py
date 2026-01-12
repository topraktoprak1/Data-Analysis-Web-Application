from openpyxl import load_workbook
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.io as pio
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import json
from PIL import Image
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode, DataReturnMode, JsCode

st.set_page_config(page_title="Excel Data Analyzer", layout="wide")

# Cache data loading for better performance
@st.cache_data(show_spinner="Loading Excel file...")
def load_excel_data(file):
    """Load Excel file and return DataFrame from DATABASE sheet if it exists"""
    # Get file extension
    file_name = file.name if hasattr(file, 'name') else ''
    
    # For xlsb files, use pyxlsb engine and try to load DATABASE sheet
    if file_name.lower().endswith('.xlsb'):
        try:
            # Try to read DATABASE sheet specifically
            df = pd.read_excel(file, sheet_name='DATABASE', engine='pyxlsb')
            return df
        except:
            # If DATABASE sheet not found, read first sheet
            df = pd.read_excel(file, engine='pyxlsb')
            return df
    else:
        # For xlsx/xls files, try DATABASE sheet first
        try:
            df = pd.read_excel(file, sheet_name='DATABASE')
            return df
        except:
            # If DATABASE sheet not found, read first sheet
            df = pd.read_excel(file)
            return df

# Cache calculated columns creation
@st.cache_data(show_spinner="Processing columns...")
def add_calculated_columns(df):
    """Add KAR/ZARAR and BF KAR/ZARAR columns, and format (Week / Month) column"""
    df = df.copy()
    
    col_isveren = None
    col_general = None
    col_birim = None
    col_hourly = None
    
    # Find columns (handle potential variations in column names)
    for col in df.columns:
        col_clean = str(col).strip()
        if 'İşveren- Hakediş (USD)' in col_clean or 'İşveren-Hakediş (USD)' in col_clean:
            col_isveren = col
        elif 'General Total' in col_clean and 'Cost (USD)' in col_clean:
            col_general = col
        elif 'İşveren-Hakediş Birim Fiyat' in col_clean and '(USD)' in col_clean:
            col_birim = col
        elif 'Hourly Unit Rate (USD)' in col_clean:
            col_hourly = col
    
    # Create KAR/ZARAR column
    if col_isveren and col_general:
        df['KAR/ZARAR'] = pd.to_numeric(df[col_isveren], errors='coerce') - pd.to_numeric(df[col_general], errors='coerce')
    
    # Create BF KAR/ZARAR column
    if col_birim and col_hourly:
        df['BF KAR/ZARAR'] = pd.to_numeric(df[col_birim], errors='coerce') - pd.to_numeric(df[col_hourly], errors='coerce')
    
    # Format (Week / Month) column if exists
    week_month_col = None
    for col in df.columns:
        if 'week' in str(col).lower() and 'month' in str(col).lower():
            week_month_col = col
            break
    
    if week_month_col:
        def format_week_month(val):
            try:
                if pd.isna(val):
                    return val
                # Try to parse as datetime
                dt = pd.to_datetime(val, errors='coerce')
                if pd.notna(dt):
                    # Format as "17/Nov"
                    return dt.strftime('%d/%b')
                return val
            except:
                return val
        
        df[week_month_col] = df[week_month_col].apply(format_week_month)
    
    return df

# Absolute path for favorite_reports.json - always save in the same location as app.py
FAVORITES_FILE_PATH = r'c:\Users\refik.toprak\Desktop\Analiz_Programı\favorite_reports.json'

# Load/Save favorite reports from/to file
def load_favorites():
    try:
        with open(FAVORITES_FILE_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)
            # Handle both old format (just dict) and new format (dict with metadata)
            if isinstance(data, dict) and '_last_loaded' in data:
                return data
            else:
                return {'reports': data, '_last_loaded': None}
    except FileNotFoundError:
        return {'reports': {}, '_last_loaded': None}

def save_favorites(favorites, last_loaded=None):
    data = {
        'reports': favorites,
        '_last_loaded': last_loaded
    }
    with open(FAVORITES_FILE_PATH, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)

# Initialize session state for favorite reports
if 'favorite_reports' not in st.session_state:
    data = load_favorites()
    st.session_state.favorite_reports = data.get('reports', {})
    st.session_state.last_loaded_report = data.get('_last_loaded', None)

if 'current_report_config' not in st.session_state:
    st.session_state.current_report_config = None

st.title("📊 Excel Data Analyzer")

# File upload
uploaded_file = st.file_uploader("Upload your Excel file", type=['xlsx', 'xls', 'xlsb'], key='file_uploader')

if uploaded_file is not None:
    try:
        # Load data with caching
        df = load_excel_data(uploaded_file)
        
        # Validate dataframe
        if df is None or df.empty:
            st.error("The uploaded file appears to be empty or could not be read.")
            st.stop()
        
        df = add_calculated_columns(df)
        
        st.success(f"File uploaded successfully! Shape: {df.shape}")
    except Exception as e:
        st.error(f"Error loading file: {str(e)}")
        st.info("Please ensure the file is a valid Excel file and not corrupted.")
        st.stop()
    
    # Sidebar for display options
    st.sidebar.header("📌 Options")
    
    # Favorite Reports Section
    st.sidebar.markdown("---")
    st.sidebar.subheader("⭐ Favorite Reports")
    
    # Auto-load last report if not already loaded
    if st.session_state.last_loaded_report and not st.session_state.current_report_config:
        if st.session_state.last_loaded_report in st.session_state.favorite_reports:
            st.session_state.current_report_config = st.session_state.favorite_reports[st.session_state.last_loaded_report]
    
    # Load favorite report dropdown
    if st.session_state.favorite_reports:
        # Set default selection to last loaded report
        report_list = ["-- Select a report --"] + list(st.session_state.favorite_reports.keys())
        default_index = 0
        if st.session_state.last_loaded_report and st.session_state.last_loaded_report in report_list:
            default_index = report_list.index(st.session_state.last_loaded_report)
        
        selected_favorite = st.sidebar.selectbox(
            "Load Saved Report",
            report_list,
            index=default_index,
            key="favorite_selector"
        )
        
        # Auto-apply if a report is selected in dropdown (even without clicking Load)
        if selected_favorite != "-- Select a report --" and not st.session_state.current_report_config:
            # Clear filter visibility state to allow loading saved values
            st.session_state.filter_visibility_state = {}
            st.session_state.current_report_config = st.session_state.favorite_reports[selected_favorite]
            st.session_state.last_loaded_report = selected_favorite
        
        if selected_favorite != "-- Select a report --":
            if st.sidebar.button("📂 Load Report"):
                # Clear filter visibility state to allow loading saved values
                st.session_state.filter_visibility_state = {}
                st.session_state.current_report_config = st.session_state.favorite_reports[selected_favorite]
                st.session_state.last_loaded_report = selected_favorite
                save_favorites(st.session_state.favorite_reports, selected_favorite)  # Save with last loaded
                st.success(f"Loaded report: {selected_favorite}")
                st.rerun()
        
        # Delete favorite report
        if selected_favorite != "-- Select a report --":
            if st.sidebar.button("🗑️ Delete Report"):
                del st.session_state.favorite_reports[selected_favorite]
                # Clear last_loaded if we're deleting it
                if st.session_state.last_loaded_report == selected_favorite:
                    st.session_state.last_loaded_report = None
                save_favorites(st.session_state.favorite_reports, st.session_state.last_loaded_report)  # Save to file
                st.success(f"Deleted report: {selected_favorite}")
                st.rerun()
    else:
        st.sidebar.info("No saved reports yet")
    
    st.sidebar.markdown("---")
    
    # Find Name Surname column first
    name_col = None
    for col in df.columns:
        col_lower = str(col).lower().replace(' ', '').replace('-', '').replace('_', '')
        if 'name' in col_lower and 'surname' in col_lower:
            name_col = col
            break
        elif col_lower in ['namesurname', 'name', 'fullname', 'adsoyad']:
            name_col = col
            break
    
    # Create filter columns list
    filter_cols = []
    for col in df.columns:
        if df[col].dtype == 'object' or not pd.api.types.is_numeric_dtype(df[col]):
            unique_values = df[col].unique()
            # Always include Name Surname, limit others to 50 unique values
            if col == name_col or len(unique_values) <= 50:
                filter_cols.append(col)
    
    # Priority: put Name Surname first
    if name_col and name_col in filter_cols:
        filter_cols.remove(name_col)
        filter_cols.insert(0, name_col)
    
    # Load saved filter controls if available
    saved_filter_controls = st.session_state.current_report_config.get('filter_controls', {}) if st.session_state.current_report_config else {}
    
    # Initialize filter visibility state
    if 'filter_visibility_state' not in st.session_state:
        st.session_state.filter_visibility_state = {}
    
    # Sidebar checkboxes for individual filter visibility
    with st.sidebar.expander("🔍 Filter Controls", expanded=True):
        filter_visibility = {}
        filter_controls_changed = False
        
        for filter_col in filter_cols:
            # Use saved value from loaded report, or from session state, or default to True
            if saved_filter_controls and filter_col in saved_filter_controls:
                default_value = saved_filter_controls[filter_col]
            else:
                default_value = st.session_state.filter_visibility_state.get(filter_col, True)
            
            filter_visibility[filter_col] = st.checkbox(
                filter_col, 
                value=default_value, 
                key=f"show_{filter_col}"
            )
            
            # Check if value changed
            old_value = st.session_state.filter_visibility_state.get(filter_col)
            if old_value is not None and old_value != filter_visibility[filter_col]:
                filter_controls_changed = True
            
            # Update session state
            st.session_state.filter_visibility_state[filter_col] = filter_visibility[filter_col]
        
        # Auto-save filter controls if they changed and a report is loaded
        if filter_controls_changed and st.session_state.last_loaded_report:
            if st.session_state.last_loaded_report in st.session_state.favorite_reports:
                # Update the filter_controls in the saved report
                st.session_state.favorite_reports[st.session_state.last_loaded_report]['filter_controls'] = filter_visibility.copy()
                # Update current config as well
                if st.session_state.current_report_config:
                    st.session_state.current_report_config['filter_controls'] = filter_visibility.copy()
                # Save to file
                save_favorites(st.session_state.favorite_reports, st.session_state.last_loaded_report)
    
    st.sidebar.markdown("---")    # Filters at the top (non-numerical columns only)
    filtered_df = df.copy()
    
    # Initialize quick_filters in session state if not exists
    if 'quick_filters' not in st.session_state:
        st.session_state.quick_filters = {}
    
    # Load filters from favorite if available
    saved_filters = st.session_state.current_report_config.get('filters', {}) if st.session_state.current_report_config else {}
    
    # Merge saved filters with quick filters
    merged_filters = {}
    for col in saved_filters:
        merged_filters[col] = saved_filters[col].copy() if isinstance(saved_filters[col], list) else [saved_filters[col]]
    
    for col in st.session_state.quick_filters:
        if col in merged_filters:
            # Merge and deduplicate
            merged_filters[col] = list(set(merged_filters[col] + st.session_state.quick_filters[col]))
        else:
            merged_filters[col] = st.session_state.quick_filters[col].copy()
    
    active_filters = {}  # Track active filter selections
    
    # Get visible filters
    visible_filters = [col for col in filter_cols if filter_visibility.get(col, True)]
    
    if visible_filters:
        col_header1, col_header2 = st.columns([3, 1])
        with col_header1:
            st.header("🔍 Filters")
        with col_header2:
            if st.session_state.quick_filters:
                if st.button("🗑️ Clear Quick Filters"):
                    st.session_state.quick_filters = {}
                    st.rerun()
        
        # Display filters in rows of 3 columns
        filters_per_row = 3
        
        with st.expander("🔽 Click to expand/collapse filters", expanded=True):
            for i in range(0, len(visible_filters), filters_per_row):
                cols = st.columns(filters_per_row)
                for j, filter_col in enumerate(visible_filters[i:i+filters_per_row]):
                    with cols[j]:
                        # Get unique values from the currently filtered dataframe (cascading)
                        unique_values = filtered_df[filter_col].dropna().unique().tolist()
                        if filter_col == name_col:
                            unique_values = sorted(unique_values)
                        
                        # Use merged filter values (saved + quick filters) if available and valid
                        saved_values = merged_filters.get(filter_col, [])
                        # Ensure saved_values is a list
                        if not isinstance(saved_values, list):
                            saved_values = []
                        # Only use saved values that exist in current data
                        # Convert both to strings for comparison to handle type mismatches
                        unique_values_str = [str(v) for v in unique_values]
                        saved_values_str = [str(v) for v in saved_values]
                        valid_saved_values = [unique_values[i] for i, v_str in enumerate(unique_values_str) if v_str in saved_values_str]
                        default_values = valid_saved_values if valid_saved_values else None
                        
                        selected = st.multiselect(f"{filter_col}", unique_values, default=default_values, key=f"filter_{filter_col}")
                        if selected:
                            filtered_df = filtered_df[filtered_df[filter_col].isin(selected)]
                            active_filters[filter_col] = selected
        
        st.markdown("---")
    
    # Sidebar display options
    st.sidebar.markdown("---")
    
    # Show raw data
    if st.sidebar.checkbox("Show Raw Data", True):
        st.subheader("📋 Raw Data")
        st.dataframe(df)
    
    # Data Info
    if st.sidebar.checkbox("Show Data Info"):
        st.subheader("ℹ️ Dataset Information")
        col1, col2, col3 = st.columns(3)
        col1.metric("Rows", df.shape[0])
        col2.metric("Columns", df.shape[1])
        col3.metric("Memory", f"{df.memory_usage(deep=True).sum() / 1024:.2f} KB")
        
        st.write("**Column Types:**")
        st.write(df.dtypes)
        
        st.write("**Missing Values:**")
        st.write(df.isnull().sum())
    
    # Statistical Summary
    if st.sidebar.checkbox("Show Statistical Summary"):
        st.subheader("📊 Statistical Summary")
        st.write(df.describe())
    
    st.subheader(f"🔍 Filtered Data ({filtered_df.shape[0]} rows)")
    st.caption("💡 Right-click on column headers or use the filter icon to filter data")
    
    # Aggregation toggle
    col_agg1, col_agg2 = st.columns([1, 4])
    with col_agg1:
        aggregate_data = st.checkbox("📊 Aggregate duplicates", value=False, help="Group by ID/Name and sum numeric columns")
    with col_agg2:
        if aggregate_data:
            # Determine grouping column
            if 'ID' in filtered_df.columns:
                group_col = st.selectbox("Group by:", ['ID', 'Name Surname'] if 'Name Surname' in filtered_df.columns else ['ID'], key='group_col')
            elif 'Name Surname' in filtered_df.columns:
                group_col = 'Name Surname'
            else:
                st.warning("No ID or Name Surname column found for aggregation")
                aggregate_data = False
                group_col = None
    
    # Apply aggregation if enabled
    if aggregate_data and group_col:
        # Columns that should NOT be summed even if numeric
        no_sum_columns = ['Control-1', 'TM Liste', 'TM Kod', 'NO-1', 'Konrol-1', 'Knrtol-2', 'NO-2', 'NO-3', 'NO-10', 'ID']
        
        # Build aggregation dictionary for ALL columns in original order
        agg_dict = {}
        for col in filtered_df.columns:
            if col == group_col:
                continue  # Skip the grouping column itself
            elif col in no_sum_columns:
                agg_dict[col] = 'first'  # Keep first value for excluded columns
            elif pd.api.types.is_numeric_dtype(filtered_df[col]):
                agg_dict[col] = 'sum'  # Sum all other numeric columns
            else:
                agg_dict[col] = 'first'  # Keep first value for text columns
        
        # Group and aggregate - this maintains column order
        display_df = filtered_df.groupby(group_col, as_index=False).agg(agg_dict)
        
        # Reorder columns to match original filtered_df
        display_df = display_df[filtered_df.columns]
        
        st.caption(f"✅ Aggregated by {group_col}: {display_df.shape[0]} unique entries (from {filtered_df.shape[0]} rows)")
    else:
        display_df = filtered_df
    
    # Configure AgGrid similar to Dash DataTable
    gb = GridOptionsBuilder.from_dataframe(display_df)
    
    # Configure each column with filtering - enable column menu
    for col in display_df.columns:
        gb.configure_column(
            col,
            filter='agTextColumnFilter',  # Text filter with search
            floatingFilter=True,  # Show filter row below headers
            sortable=True,
            resizable=True,
            suppressMenu=False,  # Enable column menu
            menuTabs=['filterMenuTab', 'generalMenuTab'],
            wrapText=False,
            autoHeight=False,
            minWidth=100,  # Minimum width
            maxWidth=300   # Maximum width
        )
    
    # Grid options with column menu enabled
    gb.configure_grid_options(
        domLayout='normal',
        enableRangeSelection=False,
        suppressMenuHide=False,  # Show menu
        animateRows=True,
        pagination=True,
        paginationPageSize=50,
        suppressContextMenu=False,
        suppressColumnVirtualisation=True,
        suppressRowVirtualisation=False
    )
    
    # Enable default column menu with auto-sizing
    gb.configure_default_column(
        filter=True,
        floatingFilter=True,
        sortable=True,
        resizable=True,
        minWidth=100,
        flex=1  # Auto-size columns to fit
    )
    
    gridOptions = gb.build()
    
    # Display AgGrid with Dash-like filtering
    grid_response = AgGrid(
        display_df,
        gridOptions=gridOptions,
        update_mode=GridUpdateMode.MODEL_CHANGED,
        data_return_mode=DataReturnMode.FILTERED_AND_SORTED,
        fit_columns_on_grid_load=True,  # Auto-fit columns on load
        theme='streamlit',
        height=600,
        allow_unsafe_jscode=True,
        enable_enterprise_modules=False,
        reload_data=False
    )
    
    # Update filtered_df based on grid filtering
    if grid_response and 'data' in grid_response:
        filtered_from_grid = pd.DataFrame(grid_response['data'])
        if len(filtered_from_grid) < len(filtered_df):
            filtered_df = filtered_from_grid
            st.info(f"Grid filtered to {len(filtered_df)} rows")
    
    # Pivot Table
    st.header("📐 Pivot Table")
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    
    # Load from favorite if available
    pivot_defaults = st.session_state.current_report_config.get('pivot', {}) if st.session_state.current_report_config else {}
    
    with col1:
        index_col = st.selectbox("Index (Rows)", ["None"] + list(filtered_df.columns), 
                                index=(["None"] + list(filtered_df.columns)).index(pivot_defaults.get('index', 'None')) if pivot_defaults.get('index', 'None') in (["None"] + list(filtered_df.columns)) else 0)
    with col2:
        columns_col = st.selectbox("Columns", ["None"] + list(filtered_df.columns),
                                  index=(["None"] + list(filtered_df.columns)).index(pivot_defaults.get('columns', 'None')) if pivot_defaults.get('columns', 'None') in (["None"] + list(filtered_df.columns)) else 0)
    with col3:
        numeric_cols = filtered_df.select_dtypes(include=['number']).columns.tolist()
        # Get default values and ensure they exist in current numeric columns
        default_values = pivot_defaults.get('values', [])
        if default_values:
            default_values = [v for v in default_values if v in numeric_cols]
        values_cols = st.multiselect("Values (select multiple)", numeric_cols, default=default_values if default_values else None)
    with col4:
        aggfunc = st.selectbox("Aggregation", ["sum", "mean", "count", "min", "max", "std"],
                              index=["sum", "mean", "count", "min", "max", "std"].index(pivot_defaults.get('aggfunc', 'sum')) if pivot_defaults.get('aggfunc', 'sum') in ["sum", "mean", "count", "min", "max", "std"] else 0)
    with col5:
        sort_col = st.selectbox("Sort By", ["None"] + (values_cols if values_cols else []))
    with col6:
        sort_order = st.selectbox("Sort Order", ["Descending", "Ascending"])
    
    if index_col != "None" and values_cols:
        try:
            # Ensure the index column and values columns have proper data types
            temp_df = filtered_df.copy()
            
            # Convert values columns to numeric if they aren't already
            for val_col in values_cols:
                if val_col in temp_df.columns:
                    temp_df[val_col] = pd.to_numeric(temp_df[val_col], errors='coerce')
            
            if columns_col != "None":
                pivot = pd.pivot_table(temp_df, values=values_cols, index=index_col, 
                                     columns=columns_col, aggfunc=aggfunc, fill_value=0, margins=True)
                # Reset index to show index as a column
                pivot = pivot.reset_index()
            else:
                # Multiple values without column grouping - already has reset_index
                pivot = temp_df.groupby(index_col)[values_cols].agg(aggfunc).reset_index()
            
            # Apply sorting if specified
            if sort_col != "None" and sort_col in pivot.columns:
                ascending = (sort_order == "Ascending")
                pivot = pivot.sort_values(by=sort_col, ascending=ascending)
            
            # Round all numeric columns to whole numbers
            numeric_columns = pivot.select_dtypes(include=['number']).columns
            for col in numeric_columns:
                pivot[col] = pivot[col].round(0).astype(int)
            
            # Apply thick borders and styling - hide index completely
            def style_pivot(df):
                return df.style.hide(axis='index').set_properties(**{
                    'border': '2px solid #000000',
                    'text-align': 'center',
                    'padding': '8px'
                }).set_table_styles([
                    {'selector': 'th', 'props': [
                        ('background-color', '#f0f0f0'),
                        ('color', '#000000'),
                        ('border', '2px solid #000000'),
                        ('text-align', 'center'),
                        ('font-weight', 'bold'),
                        ('padding', '10px')
                    ]},
                    {'selector': 'td', 'props': [
                        ('border', '2px solid #000000')
                    ]},
                    {'selector': '', 'props': [
                        ('border-collapse', 'collapse')
                    ]}
                ])
            
            styled_pivot = style_pivot(pivot)
            
            # Make the pivot table scrollable with fixed height
            st.markdown("""
                <style>
                .pivot-container {
                    max-height: 500px;
                    overflow: auto;
                    border: 3px solid #000000;
                    border-radius: 5px;
                }
                .pivot-container table {
                    width: 100%;
                }
                </style>
            """, unsafe_allow_html=True)
            
            st.markdown(f'<div class="pivot-container">{styled_pivot.to_html(index=False)}</div>', unsafe_allow_html=True)
            
            # Download pivot table
            csv = pivot.to_csv(index=False).encode('utf-8')
            st.download_button("Download Pivot Table", csv, "pivot_table.csv", "text/csv")
        except Exception as e:
            st.error(f"Error creating pivot table: {str(e)}")
    
    # Visualizations
    st.header("📈 Visualizations")
    
    # Load saved chart settings if available
    saved_charts = []
    if st.session_state.current_report_config and 'charts' in st.session_state.current_report_config:
        saved_charts = st.session_state.current_report_config['charts']
    
    # Number of charts selector
    default_num_charts = len(saved_charts) if saved_charts else 1
    num_charts = st.number_input("Number of Charts to Display", min_value=1, max_value=6, value=default_num_charts, step=1)
    st.markdown("---")
    
    # Get numeric columns for chart options
    numeric_cols = filtered_df.select_dtypes(include=['number']).columns.tolist()
    
    # Store chart figures for Word export
    chart_figures = []
    chart_titles = []
    chart_configs = []  # Track chart configurations
    
    # Create charts in columns (2 per row for better layout)
    charts_per_row = 2
    for chart_idx in range(num_charts):
        # Start a new row every 2 charts
        if chart_idx % charts_per_row == 0:
            cols = st.columns(charts_per_row)
        
        with cols[chart_idx % charts_per_row]:
            st.subheader(f"Chart {chart_idx + 1}")
            
            # Get saved config for this chart if available
            saved_chart = saved_charts[chart_idx] if chart_idx < len(saved_charts) else {}
            default_viz = saved_chart.get('type', 'Bar Chart')
            
            viz_type = st.selectbox("Select Visualization Type", 
                           ["Bar Chart", "Line Chart", "Scatter Plot", "Pie Chart", 
                            "Box Plot", "Histogram", "Heatmap"], 
                           index=["Bar Chart", "Line Chart", "Scatter Plot", "Pie Chart", 
                                  "Box Plot", "Histogram", "Heatmap"].index(default_viz) if default_viz in ["Bar Chart", "Line Chart", "Scatter Plot", "Pie Chart", "Box Plot", "Histogram", "Heatmap"] else 0,
                           key=f"viz_type_{chart_idx}")
            
            if viz_type == "Bar Chart":
                default_x = saved_chart.get('x_axis', filtered_df.columns[0])
                default_y = saved_chart.get('y_axis', numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                default_color = saved_chart.get('color_by', 'None')
                
                x_axis = st.selectbox("X-axis", filtered_df.columns, 
                                     index=list(filtered_df.columns).index(default_x) if default_x in filtered_df.columns else 0,
                                     key=f"bar_x_{chart_idx}")
                y_axis = st.selectbox("Y-axis", numeric_cols, 
                                     index=numeric_cols.index(default_y) if default_y in numeric_cols else 0,
                                     key=f"bar_y_{chart_idx}")
                color_by = st.selectbox("Color by", ["None"] + list(filtered_df.columns), 
                                       index=(["None"] + list(filtered_df.columns)).index(default_color) if default_color in ["None"] + list(filtered_df.columns) else 0,
                                       key=f"bar_color_{chart_idx}")
        
                # Aggregate data by X-axis
                if color_by != "None":
                    agg_df = filtered_df.groupby([x_axis, color_by], as_index=False)[y_axis].sum()
                    fig = px.bar(agg_df, x=x_axis, y=y_axis, color=color_by,
                                title=f"{y_axis} by {x_axis}", barmode='group')
                else:
                    agg_df = filtered_df.groupby(x_axis, as_index=False)[y_axis].sum()
                    fig = px.bar(agg_df, x=x_axis, y=y_axis, title=f"{y_axis} by {x_axis}")
                
                fig.update_layout(
                    xaxis_tickangle=-45,
                    xaxis_title=x_axis,
                    yaxis_title=y_axis,
                    showlegend=True if color_by != "None" else False,
                    height=400
                )
                fig.update_traces(texttemplate='%{y:.2s}', textposition='outside')
                st.plotly_chart(fig, width='stretch', key=f"bar_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append(f"Bar Chart: {y_axis} by {x_axis}")
                chart_configs.append({'type': 'Bar Chart', 'x_axis': x_axis, 'y_axis': y_axis, 'color_by': color_by})
    
            elif viz_type == "Line Chart":
                default_x = saved_chart.get('x_axis', filtered_df.columns[0])
                default_y = saved_chart.get('y_axis', numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                default_color = saved_chart.get('color_by', 'None')
                
                x_axis = st.selectbox("X-axis", filtered_df.columns, 
                                     index=list(filtered_df.columns).index(default_x) if default_x in filtered_df.columns else 0,
                                     key=f"line_x_{chart_idx}")
                y_axis = st.selectbox("Y-axis", numeric_cols, 
                                     index=numeric_cols.index(default_y) if default_y in numeric_cols else 0,
                                     key=f"line_y_{chart_idx}")
                color_by = st.selectbox("Color by", ["None"] + list(filtered_df.columns), 
                                       index=(["None"] + list(filtered_df.columns)).index(default_color) if default_color in ["None"] + list(filtered_df.columns) else 0,
                                       key=f"line_color_{chart_idx}")
        
                if color_by != "None":
                    agg_df = filtered_df.groupby([x_axis, color_by], as_index=False)[y_axis].sum()
                    fig = px.line(agg_df, x=x_axis, y=y_axis, color=color_by,
                                 title=f"{y_axis} over {x_axis}", markers=True)
                else:
                    agg_df = filtered_df.groupby(x_axis, as_index=False)[y_axis].sum()
                    fig = px.line(agg_df, x=x_axis, y=y_axis, 
                                 title=f"{y_axis} over {x_axis}", markers=True)
                
                fig.update_layout(
                    xaxis_tickangle=-45,
                    xaxis_title=x_axis,
                    yaxis_title=y_axis,
                    height=400,
                    hovermode='x unified'
                )
                fig.update_traces(line=dict(width=3))
                st.plotly_chart(fig, width='stretch', key=f"line_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append(f"Line Chart: {y_axis} over {x_axis}")
                chart_configs.append({'type': 'Line Chart', 'x_axis': x_axis, 'y_axis': y_axis, 'color_by': color_by})
    
            elif viz_type == "Scatter Plot":
                default_x = saved_chart.get('x_axis', numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                default_y = saved_chart.get('y_axis', numeric_cols[1] if len(numeric_cols) > 1 else numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                default_color = saved_chart.get('color_by', 'None')
                default_size = saved_chart.get('size_by', 'None')
                
                x_axis = st.selectbox("X-axis", numeric_cols, 
                                     index=numeric_cols.index(default_x) if default_x in numeric_cols else 0,
                                     key=f"scatter_x_{chart_idx}")
                y_axis = st.selectbox("Y-axis", numeric_cols, 
                                     index=numeric_cols.index(default_y) if default_y in numeric_cols else 0,
                                     key=f"scatter_y_{chart_idx}")
                color = st.selectbox("Color by", ["None"] + list(filtered_df.columns), 
                                    index=(["None"] + list(filtered_df.columns)).index(default_color) if default_color in ["None"] + list(filtered_df.columns) else 0,
                                    key=f"scatter_color_{chart_idx}")
                size = st.selectbox("Size by", ["None"] + numeric_cols, 
                                   index=(["None"] + numeric_cols).index(default_size) if default_size in ["None"] + numeric_cols else 0,
                                   key=f"scatter_size_{chart_idx}")
        
                color_val = None if color == "None" else color
                size_val = None if size == "None" else size
                
                fig = px.scatter(filtered_df, x=x_axis, y=y_axis, color=color_val, size=size_val,
                                title=f"{y_axis} vs {x_axis}",
                                hover_data=filtered_df.columns)
                
                fig.update_layout(
                    xaxis_title=x_axis,
                    yaxis_title=y_axis,
                    height=400
                )
                fig.update_traces(marker=dict(opacity=0.7, line=dict(width=0.5, color='white')))
                st.plotly_chart(fig, width='stretch', key=f"scatter_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append(f"Scatter Plot: {y_axis} vs {x_axis}")
                chart_configs.append({'type': 'Scatter Plot', 'x_axis': x_axis, 'y_axis': y_axis, 'color_by': color, 'size_by': size})
    
            elif viz_type == "Pie Chart":
                default_names = saved_chart.get('names', filtered_df.columns[0])
                default_values = saved_chart.get('values', numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                
                names = st.selectbox("Labels", filtered_df.columns, 
                                    index=list(filtered_df.columns).index(default_names) if default_names in filtered_df.columns else 0,
                                    key=f"pie_names_{chart_idx}")
                values = st.selectbox("Values", numeric_cols, 
                                     index=numeric_cols.index(default_values) if default_values in numeric_cols else 0,
                                     key=f"pie_values_{chart_idx}")
        
                agg_df = filtered_df.groupby(names, as_index=False)[values].sum()
                agg_df = agg_df.nlargest(10, values)
                
                fig = px.pie(agg_df, names=names, values=values, 
                            title=f"{values} by {names} (Top 10)",
                            hole=0.3)
                
                fig.update_traces(
                    textposition='outside',
                    textinfo='label+percent',
                    marker=dict(line=dict(color='white', width=2))
                )
                fig.update_layout(height=400)
                st.plotly_chart(fig, width='stretch', key=f"pie_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append(f"Pie Chart: {values} by {names}")
                chart_configs.append({'type': 'Pie Chart', 'names': names, 'values': values})
    
            elif viz_type == "Box Plot":
                default_x = saved_chart.get('x_axis', filtered_df.columns[0])
                default_y = saved_chart.get('y_axis', numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                default_color = saved_chart.get('color_by', 'None')
                
                x_axis = st.selectbox("Category", filtered_df.columns, 
                                     index=list(filtered_df.columns).index(default_x) if default_x in filtered_df.columns else 0,
                                     key=f"box_x_{chart_idx}")
                y_axis = st.selectbox("Values", numeric_cols, 
                                     index=numeric_cols.index(default_y) if default_y in numeric_cols else 0,
                                     key=f"box_y_{chart_idx}")
                color_by = st.selectbox("Color by", ["None"] + list(filtered_df.columns), 
                                       index=(["None"] + list(filtered_df.columns)).index(default_color) if default_color in ["None"] + list(filtered_df.columns) else 0,
                                       key=f"box_color_{chart_idx}")
        
                color_val = None if color_by == "None" else color_by
                
                fig = px.box(filtered_df, x=x_axis, y=y_axis, color=color_val,
                            title=f"{y_axis} Distribution by {x_axis}",
                            points="outliers")
                
                fig.update_layout(
                    xaxis_tickangle=-45,
                    xaxis_title=x_axis,
                    yaxis_title=y_axis,
                    height=400
                )
                fig.update_traces(marker=dict(opacity=0.6))
                st.plotly_chart(fig, width='stretch', key=f"box_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append(f"Box Plot: {y_axis} by {x_axis}")
                chart_configs.append({'type': 'Box Plot', 'x_axis': x_axis, 'y_axis': y_axis, 'color_by': color_by})
    
            elif viz_type == "Histogram":
                default_column = saved_chart.get('column', numeric_cols[0] if numeric_cols else filtered_df.columns[0])
                default_bins = saved_chart.get('bins', 30)
                default_color = saved_chart.get('color_by', 'None')
                
                column = st.selectbox("Select Column", numeric_cols, 
                                     index=numeric_cols.index(default_column) if default_column in numeric_cols else 0,
                                     key=f"hist_{chart_idx}")
                bins = st.slider("Number of Bins", 10, 100, default_bins, key=f"bins_{chart_idx}")
                color_by = st.selectbox("Color by", ["None"] + list(filtered_df.columns), 
                                       index=(["None"] + list(filtered_df.columns)).index(default_color) if default_color in ["None"] + list(filtered_df.columns) else 0,
                                       key=f"hist_color_{chart_idx}")
        
                color_val = None if color_by == "None" else color_by
                
                fig = px.histogram(filtered_df, x=column, nbins=bins, color=color_val,
                                  title=f"Distribution of {column}",
                                  marginal="box")
                
                fig.update_layout(
                    xaxis_title=column,
                    yaxis_title="Count",
                    height=400,
                    bargap=0.1
                )
                fig.update_traces(marker=dict(line=dict(width=0.5, color='white')))
                st.plotly_chart(fig, width='stretch', key=f"hist_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append(f"Histogram: {column}")
                chart_configs.append({'type': 'Histogram', 'column': column, 'bins': bins, 'color_by': color_by})
    
            elif viz_type == "Heatmap":
                corr_df = filtered_df.select_dtypes(include=['number']).corr()
                fig = px.imshow(corr_df, text_auto=True, aspect="auto", 
                               title="Correlation Heatmap")
                fig.update_layout(height=400)
                st.plotly_chart(fig, width='stretch', key=f"heatmap_chart_{chart_idx}")
                
                # Store for Word export
                chart_figures.append(fig)
                chart_titles.append("Correlation Heatmap")
                chart_configs.append({'type': 'Heatmap'})
        
        st.markdown("---")  # Separator between charts
    
    # Save to Favorites Section
    st.markdown("---")
    st.header("⭐ Save Current Configuration to Favorites")
    
    col_save1, col_save2 = st.columns([3, 1])
    with col_save1:
        report_name = st.text_input("Report Name", placeholder="e.g., Discipline Analysis")
    with col_save2:
        st.write("")  # Spacing
        st.write("")  # Spacing
        if st.button("💾 Save to Favorites"):
            if report_name:
                # Save current configuration including filters
                st.session_state.favorite_reports[report_name] = {
                    'pivot': {
                        'index': index_col,
                        'columns': columns_col,
                        'values': values_cols,
                        'aggfunc': aggfunc
                    },
                    'filters': active_filters,  # Save active filters
                    'filter_controls': filter_visibility,  # Save filter controls (visibility)
                    'charts': chart_configs  # Save chart configurations
                }
                save_favorites(st.session_state.favorite_reports, st.session_state.last_loaded_report)  # Save to file
                st.success(f"✅ Report '{report_name}' saved to favorites!")
                st.session_state.current_report_config = None  # Clear the loaded config
            else:
                st.error("Please enter a report name")
    
    # Export Report
    st.markdown("---")
    st.header("📄 Generate Report")

    export_format = st.radio("Select export format:", ["Word (DOCX)", "Excel (XLSX) - Better for colored charts"], horizontal=True)

    if st.button("Generate Report"):
        try:
            if export_format == "Word (DOCX)":
                # Create Word document
                doc = Document()

                # Add title
                title = doc.add_heading('Data Analysis Report', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add date
                date_para = doc.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
                date_run.font.size = Pt(11)

                doc.add_paragraph()  # Spacing

                # Add dataset summary
                doc.add_heading('Dataset Overview', 1)
                doc.add_paragraph(f'Total rows: {filtered_df.shape[0]}')
                doc.add_paragraph(f'Total columns: {filtered_df.shape[1]}')
                doc.add_paragraph()

                # Add pivot table if exists
                if index_col != "None" and values_cols:
                    doc.add_heading('Pivot Table Analysis', 1)

                    # Convert pivot table to Word table
                    table = doc.add_table(rows=1, cols=len(pivot.columns))
                    table.style = 'Light Grid Accent 1'

                    # Header row
                    header_cells = table.rows[0].cells
                    for idx, col_name in enumerate(pivot.columns):
                        header_cells[idx].text = str(col_name)
                        # Bold header
                        for paragraph in header_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                    # Data rows
                    for _, row in pivot.iterrows():
                        row_cells = table.add_row().cells
                        for idx, value in enumerate(row):
                            row_cells[idx].text = str(value)

                    doc.add_paragraph()

                # Add charts
                doc.add_heading('Visualizations', 1)
                doc.add_paragraph()

                # Add each chart as an image
                if chart_figures:
                    for idx, (fig, title) in enumerate(zip(chart_figures, chart_titles)):
                        try:
                            doc.add_heading(f'Chart {idx + 1}: {title}', 2)

                            # Convert plotly figure to image using kaleido
                            img_bytes = fig.to_image(format="png", width=1400, height=800, engine="kaleido")
                            img_stream = io.BytesIO(img_bytes)

                            # Add image to document
                            doc.add_picture(img_stream, width=Inches(6))

                            # Center the image
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            doc.add_paragraph()  # Add spacing

                        except Exception as chart_error:
                            doc.add_paragraph(f'Chart: {title}')
                            doc.add_paragraph(f'(Error exporting chart: {str(chart_error)})')
                            doc.add_paragraph()
                else:
                    doc.add_paragraph('No charts to export.')
                    doc.add_paragraph()

                # Save document to BytesIO
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                # Download button
                st.download_button(
                    label="📥 Download Word Report",
                    data=doc_io,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.success("✅ Report generated successfully!")

            else:
                # Excel format
                # Create Excel workbook
                excel_io = io.BytesIO()

                with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
                    # Write filtered data
                    filtered_df.to_excel(writer, sheet_name='Data', index=False)

                    # Write pivot table if exists
                    if index_col != "None" and values_cols:
                        pivot.to_excel(writer, sheet_name='Pivot Table', index=False)

                    # Get workbook and create charts sheet
                    workbook = writer.book
                    charts_sheet = workbook.add_worksheet('Charts')

                    # Add charts as images
                    if chart_figures:
                        row_position = 0
                        for idx, (fig, title) in enumerate(zip(chart_figures, chart_titles)):
                            try:
                                # Convert plotly to image (in-memory)
                                img_bytes = fig.to_image(format="png", width=1200, height=600, engine="kaleido")
                                img_buffer = io.BytesIO(img_bytes)

                                # Insert image into Excel using in-memory bytes
                                charts_sheet.write(row_position, 0, f'Chart {idx + 1}: {title}')
                                charts_sheet.insert_image(row_position + 1, 0, 'chart.png', {'image_data': img_buffer, 'x_scale': 0.5, 'y_scale': 0.5})

                                row_position += 35  # Move to next chart position

                            except Exception as chart_error:
                                charts_sheet.write(row_position, 0, f'Chart {idx + 1}: {title} - Error: {str(chart_error)}')
                                row_position += 2

                excel_io.seek(0)

                # Download button
                st.download_button(
                    label="📥 Download Excel Report",
                    data=excel_io,
                    file_name=f"analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

                st.success("✅ Excel report generated successfully!")

        except Exception as e:
            st.error(f"Error generating report: {str(e)}")
else:
    st.info("👆 Please upload an Excel file to get started")
    st.markdown("""
    ### Features:
    - 📊 Data analysis and statistics
    - 🔍 Interactive filters
    - 📐 Customizable pivot tables
    - 📈 Multiple chart types
    - 💾 Export capabilities
    """)




